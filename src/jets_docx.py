# -*- coding: utf-8 -*-
"""
@project: xunfei-audio2text-20200723
@Author : Jet Li
@File : newfile1.py
@Time : 2020-07-24
@Desc:
Jet自定义，转换成 Word文档的工具


"""

import docx
'''
Jet: 不能安装pip install docx,否则运行时会提示错误：
  File "C:\lib\site-packages\docx.py", line 30, in <module>
    from exceptions import PendingDeprecationWarning
ModuleNotFoundError: No module named 'exceptions'
 if already installed: pip unistall docx
pip install python-docx
'''
import re  # regular expression, 字符串替换时要用到

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#from docx.opc.oxml import qn
from docx.oxml.ns import qn  # 设置中文字体
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    # element.set(docx.oxml.ns.qn(name), value)
    element.set(qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = '第 '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    # t2.text = '//' # Jet: 用转义反斜杠，隔开页码， 会出现//
    t2.text = r'/'  # Jet: 用转义反斜杠，隔开页码

    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = ' 页'
    page_run._r.append(t1)


# 添加指定格式的段落/文字
# text：段落内容
def add_formatted_text(document, file_name, text, font_cn_name="宋体", size=12,
                       color=RGBColor(0, 0, 0), alignment="left",
                       is_bold=False):
    # 添加段落
    p1 = document.add_paragraph()
    text1 = p1.add_run(text)

    text1.font.name = 'Times New Roman'  # 控制是西文时的字体
    text1.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn_name)  # 控制是中文时的字体

    text1.font.size = Pt(size)
    text1.font.color.rgb = color

    # 此处高亮显示“安静”为绿色， 如果有的话
    # 此处还需改进。。。教练的文字，字体不对，颜色蓝色也没有了
    # highlight_keywords(p1, text, keyword, color)

    # 加粗
    text1.font.bold = is_bold
    # 斜体
    # text1.font.italic = True

    if alignment == "center":
        p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif alignment == "left":
        p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p1.paragraph_format.space_after = Pt(10)  # 段后间距，default

    document.save(file_name)

# 设置doc文档中指定关键词的颜色，高亮显示：
#应用举例：icf对话中[安静]
def highlight_keywords(paragraph, content, keyword, color=RGBColor(0, 0, 0)):
    if content.find(keyword) != -1:  # 如果该行有关键字,就以关键字为分界进行分割
        pt = keyword
        res = re.split(pt, content)  # 分割

        # add_run在同一段添加内容
        run = paragraph.add_run(res[0])  # 输入关键字之前的字符
        run.font.color.rgb = color  # 字体颜色
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # print(res[0],end=' ')

        run = paragraph.add_run(keyword)  # 输入关键字
        run.font.name = u'黑体'  # 设置插入的字体
        run.font.size = Pt(12)
        # 设置关键字之后也的字体颜色，这里设置为绿色
        run.font.color.rgb = RGBColor(0, 250, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # print(res[1],end=' ')
        run = paragraph.add_run(res[1])  # 输入关键字之后的字符
        run.font.color.rgb = color  # 字体颜色
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    else:
        run = paragraph.add_run(content)  # 如果该行没有关键字，则直接输入word
        # print(content)
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(15)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 保存内容到指定Word文件
def save_content2word(filename, content):
    document = docx.Document()
    add_formatted_text(filename, document, content)
    print("保存数据为Word文件%s成功" % filename)
