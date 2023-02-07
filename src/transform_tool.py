# 从转写好的数据文件读取信息，转换成教练、客户对话的格式
# 包括时间戳
# 再用word模板，生成word文件
# start from 2020-04
import json  # json字符串转换为列表类型时要用到
import datetime  # 没用到
import re  # regular expression, 字符串替换时要用到
from collections import OrderedDict  # 导入collections.OrderedDict
import pickle  # 序列化（保存）对象到文件

# 以下生成word文档需要
from docx.oxml import OxmlElement

import jets_docx
from JetsXunfeiApi import XunfeiAudioConvertRequestApi

'''
Jet: 不能安装pip install docx,否则运行时会提示错误：
  File "C:\lib\site-packages\docx.py", line 30, in <module>
    from exceptions import PendingDeprecationWarning
ModuleNotFoundError: No module named 'exceptions'
 if already installed: pip unistall docx
pip install python-docx
'''
import docx
from docx.shared import Inches, Cm, Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 字体设置

# 修改word关键词颜色
import os
from docx.oxml.ns import qn  # 设置中文字体

# 统计字符数
import string
from collections import namedtuple

# audio_file = "./audio/2020-06-02-Sherry-facilitator.mp3"

PAUSE_MIN_SECONDS = 7  # 安静7秒以上才显示， 2020-10-08 从3秒改为7秒

speaker1 = 'speaker1'  # 可GUI获取，修改值为教练/客户
speaker2 = 'speaker2'  # 可GUI获取，修改值为教练/客户

total_time = 0  # 音频总时长
total_length_speaker1 = 0
total_length_speaker2 = 0
file_content = ''  # 返回的全部对话文本内容（包括教练和客户），之后写入指定文件
speaker1_content = ''  # 仅保存speaker1的对话
speaker2_content = ''  # 仅保存speaker2的对话
speaker1_time = 0  # 讲话时间秒数
speaker2_time = 0  # 讲话时间秒数
speaker1_quiet_time = 0  # 安静时间秒数
speaker2_quiet_time = 0  # 安静时间秒数
keyword = "安静"  # 用于高亮显示的关键词
raw_dialog_list = []  # 带有时间戳的对话列表，以便后续对照编辑


# 转换成 分:秒格式,例如：01:12
def time_format(miniseconds):
    seconds = miniseconds / 1000
    # result = datetime.timedelta(seconds)#只保留到秒数
    m, s = divmod(seconds, 60)  # 利用元组进行整除、余数计算
    h, m = divmod(m, 60)  # 小时，可根据需要添加
    if h > 0:
        result = "%02d:%02d:%02d" % (h, m, s)  # 有的音频可能超过一小时，所以还是加上小时数
    else:
        result = "%02d:%02d" % (m, s)
    # print("time:", result)
    return result


# 批量替换，供edit_common_error调用
def multiple_replace(text, custom_dict):
    '''
    re.sub(`pattern`, `repl`, `string`, `count=0`, `flags=0`)

    `pattern`, `repl`, `string` 为必选参数
    `count`, `flags` 为可选参数
    `pattern`正则表达式
    `repl`被替换的内容，可以是字符串，也可以是函数
    `string`正则表达式匹配的内容
    `count`由于正则表达式匹配的结果是多个，使用count来限定替换的个数从左向右，默认值是0，替换所有的匹配到的结果
    `flags`是匹配模式，`re.I`忽略大小写，`re.L`表示特殊字符集\w,\W,\b,\B,\s,\S，`re.M`表示多行模式，`re.S` ‘.’包括换行符在内的任意字符，`re.U`表示特殊字符集\w,\W,\b,\B,\d,\D,\s,\D
    '''
    rx = re.compile('|'.join(map(re.escape, custom_dict)))

    def one_xlat(match):
        return custom_dict[match.group(0)]

    return rx.sub(one_xlat, text)


# Jet adds: 替换讯飞语音转写的常见错误或不当标点
def edit_common_error(text):
    # 此处需要根据常见问题，添加到字典里...
    error_dict = {
        "。呢": "呢？",
        "？呢": "呢？",
        "，呢": "呢，",
        "？啊": "啊？",
        "。啊": "啊。",
        "，啊": "啊，",
        "。吧": "吧。",
        "，吧": "吧，",
        "，哈": "哈，",
        "？呀": "呀？",
        "好，呀": "好呀，",
        "，哈": "哈，",
        "。哈": "哈。",
        "很好。": "很好!",
        "非常好。": "非常好!",
        "。唉": "唉。",
        # "嗯":"嗯，",
        "，嘛": "嘛，",
        "。嘛": "嘛。"
    }
    '''
    if text == "嗯":  # 一句话只有这一个字
        text = "嗯。"
    elif "嗯" in text and len(text) > 1:
        text = text.replace("嗯", "嗯，")  # 小心检查，是否会出错？在句尾会多逗号！
    else:
        text = multiple_replace(text, error_dict)  # 调用批量替换
    '''
    text = multiple_replace(text, error_dict)  # 调用批量替换

    return text


# 统计字符数（只包括中文, 不包括英文字母、数字，由于是说话，不包括标点符号）
# Jet：这比用len计算要准确得多！且教练说话占比会明显减少。 每小时计数要减少3000字！！！
# Word计算时把中文标点当作汉字了。
# need to
# import string
def count_zh_chars(s):
    count_en = count_dg = count_sp = count_zh = count_pu = 0
    s_len = len(s)
    for c in s:
        if c in string.ascii_letters:
            count_en += 1
        elif c.isdigit():
            count_dg += 1
        elif c.isspace():
            count_sp += 1
        elif c.isalpha():
            count_zh += 1
        else:
            count_pu += 1
    # total_chars = count_zh + count_en + count_sp + count_dg + count_pu
    # total_chars = count_zh + count_en +  count_dg
    total_chars = count_zh

    return total_chars


# 向指定的Word文件 以指定格式，写入指定内容
#
def add_formatted_text(file_name, document, text, font_cn_name="宋体", size=12,
                       color=RGBColor(0, 0, 0), alignment="left",
                       is_bold=False):
    # 添加段落
    p1 = document.add_paragraph()
    text1 = p1.add_run(text)

    text1.font.name = 'Times New Roman'  # 控制是西文时的字体
    text1.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn_name)  # 控制是中文时的字体

    text1.font.size = Pt(size)
    text1.font.color.rgb = color

    # 加粗
    text1.font.bold = is_bold
    # 斜体
    # text1.font.italic = True

    if alignment == "center":
        p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif alignment == "left":
        p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p1.paragraph_format.space_after = Pt(10)  # 段后间距，default

    document.save(file_name)


"""
以下内容转移到了jets_docx.py

# 保存内容到指定Word文件
def save_content2word(filename, content):
    document = docx.Document()
    add_formatted_text(filename, document, content)
    print("保存数据为Word文件%s成功" % filename)


# 设置doc中指定关键词的颜色，高亮显示：
def highlight_keywords(paragraph, content, keyword, color=RGBColor(0, 0, 0)):
    if content.find(keyword) != -1:  # 如果该行有关键字,就以关键字为分界进行分割
        pt = keyword
        res = re.split(pt, content)  # 分割

        # add_run在同一段添加内容
        run = paragraph.add_run(res[0])  # 输入关键字之前的字符
        run.font.color.rgb = color  # 字体颜色
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # print(res[0],end=' ')

        run = paragraph.add_run(keyword)  # 输入关键字
        run.font.name = u'黑体'  # 设置插入的字体
        run.font.size = Pt(12)
        # 设置关键字之后也的字体颜色，这里设置为绿色
        run.font.color.rgb = RGBColor(0, 250, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # print(res[1],end=' ')
        run = paragraph.add_run(res[1])  # 输入关键字之后的字符
        run.font.color.rgb = color  # 字体颜色
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    else:
        run = paragraph.add_run(content)  # 如果该行没有关键字，则直接输入word
        # print(content)
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(15)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

"""


# 转换只有一位讲话者的原始转写文件
# 默认按1000ms 区分段落（添加换行）
def convert_1speaker_timestamp_rawtxt(raw_txt_file, gap_limit=1000):
    """

    :param raw_txt_file:
    :param gap_limit:
    :return:
        timestamp_list = []  # 包含起止时间戳,list里包含的字典结构如下：
        all_content = ""  # 包含所有讲话内容，没有时间戳,没有换行符
        combine_list = []  # 合并在gap_time以内的句子到一段（可能不太准确，尤其是录制课程时说一句话可能有几次断开）

    原始数据中，间隔时间小于gap_limit 毫秒数，则进行合并
    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "text": "what the speaker says"
    }

    """

    raw_list = get_result_list(raw_txt_file)  # 获取讯飞原始转写的数据信息

    timestamp_list = []  # 包含起止时间戳,list里包含的字典结构如下：
    all_content = ""  # 包含所有讲话内容，没有时间戳,没有换行符
    combine_list = []  # 合并在gap_time以内的句子到一段（可能不太准确，尤其是录制课程时说一句话可能有几次断开）

    length = len(raw_list)
    index = 0
    while index < length:  # Fix following bug:
        # for index in range(length): #用这个有Bug, 会重复提交一些记录
        item = raw_list[index]
        if item['speaker'] == '2':
            print("原始转写文件是二人对话，不适合调用1人讲话模式进行转写。")
            return

        timestamp_dict = {}  # 这里每次都得清空一次
        # timestamp_dict["speak_time"] = item['bg'] + '-' + item['ed'] # for Debug
        timestamp_dict["speak_time"] = time_format(int(item['bg'])) + '-' + time_format(int(item['ed']))

        timestamp_dict['text'] = item['onebest']
        timestamp_list.append(timestamp_dict)
        all_content = ''.join([all_content, item['onebest']])

        combine_dict = {}
        begin_time = item['bg']
        combine_text = item['onebest']
        end_time = item['ed']

        # 原始数据中，间隔时间小于gap_limit 毫秒数，则进行合并
        # to do later...
        if (index + 1) < length:
            gap_time = int(raw_list[index + 1]['bg']) - int(raw_list[index]['ed'])
            while gap_time <= gap_limit:
                index += 1
                item = raw_list[index]  # 换下一段话
                end_time = item['ed']  # 以下一段话的停止时间为这一段合并的话的停止时间

                timestamp_dict = {}  # 这里每次都得清空一次
                # timestamp_dict["speak_time"] = item['bg'] + '-' + item['ed'] # for Debug
                timestamp_dict["speak_time"] = time_format(int(item['bg'])) + '-' + time_format(int(item['ed']))
                timestamp_dict['text'] = item['onebest']
                timestamp_list.append(timestamp_dict)
                all_content = ''.join([all_content, item['onebest']])

                combine_text = ''.join([combine_text, item['onebest']])

                if (index + 1) < length:
                    gap_time = int(raw_list[index + 1]['bg']) - int(raw_list[index]['ed'])
                    continue
                else:
                    break
            # combine_dict["speak_time"] = begin_time + '-' + end_time  # for Debug
            combine_dict["speak_time"] = time_format(int(begin_time)) + '-' + time_format(int(end_time))

            combine_dict["text"] = combine_text
            combine_list.append(combine_dict)

        index += 1

    return timestamp_list, all_content, combine_list


# 按原始转写的时间戳保存到文本文件
def save_1speaker_raw_timestamp2txt(raw_txt_file, output_file):
    '''
    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "text": "what the speaker says"
    }
    '''
    result = convert_1speaker_timestamp_rawtxt(raw_txt_file)
    timestamp_list = result[0]

    try:
        fw = open(output_file, 'w', encoding='utf-8')
        text = ''
        for item in timestamp_list:
            text = ''.join([text, item['speak_time'], ' ', item['text'], '\n'])

        fw.write(text)
        print('从原始文件转成起止时间戳文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 合并原始转写的时间戳保存到文本文件，指定间隔时间秒数
def save_1speaker_combined_timestamp2txt(raw_txt_file, output_file, gap_ms=2000):
    result = convert_1speaker_timestamp_rawtxt(raw_txt_file, gap_limit=gap_ms)
    combine_list = result[2]
    try:
        fw = open(output_file, 'w', encoding='utf-8')
        text = ''
        for item in combine_list:
            text = ''.join([text, item['speak_time'], ' ', item['text'], '\n'])

        fw.write(text)
        print('从原始文件按间隔合并讲话，转成起止时间戳文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


def save_1speaker_only_txt(raw_txt_file, output_file):
    result = convert_1speaker_timestamp_rawtxt(raw_txt_file)
    all_content = result[1]

    try:
        fw = open(output_file, 'w', encoding='utf-8')
        fw.write(all_content)
        print('从原始文件转成纯文字文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 转换成列表（需要按对话顺序）
# 同一个人的讲话，合并到一起，直到换另一个讲话
# 注：如果转写结果本身把说话者搞错了，那么结果就错得厉害，并且连时间戳也会出问题！
# 需要显示给用户，由用户手工编辑更正！
# 二人对话的中间，可能有安静的时间！处理过程加入安静的时间
def convert_2speakers_rawtxt(raw_txt_file):
    global file_content  # 同一个人的讲话，合并到一起，直到换另一个讲话
    global total_time
    global total_length_speaker1
    global total_length_speaker2
    global speaker1_content
    global speaker2_content
    global speaker1_time
    global speaker2_time
    global speaker1_quiet_time
    global speaker2_quiet_time

    global raw_dialog_list  # 包含起止时间戳,list里包含的字典结构如下：
    '''
                raw_dialog_dict = {}  #2020/07/02 Fix Bug.这里每次都得清空一次
                raw_dialog_dict['begin'] = time_format(int(item['bg']))
                raw_dialog_dict['end'] = time_format(int(item['ed']))
                raw_dialog_dict['speaker'] = 'speaker1'
                raw_dialog_dict['text'] = item['onebest']

    '''


    # 初始化归零，避免重复调用时递增计算
    total_length_speaker1 = 0
    total_length_speaker2 = 0
    speaker1_time = 0
    speaker2_time = 0
    speaker1_quiet_time = 0
    speaker2_quiet_time = 0

    raw_list = get_result_list(raw_txt_file)  # 获取讯飞原始转写的数据信息
    # print('raw:', raw_list)
    conversation_list = []  # 记录所有对话，按speaker1, speaker2交替的形式存储, 没用上！
    index = 0
    length = len(raw_list)
    item_s1 = {}  # 记录每次对话speaker1说的时间戳和内容
    item_s2 = {}  # 记录每次对话speaker2说的时间戳和内容
    end = 0

    # raw_dialog_dict = {}  # 写在这里不对！；不能用OrderedDict()#后续访问需要按添加的顺序访问
    raw_dialog_list = []  # 存储原始对话数据（包括超过指定秒数的安静时长）

    # 2020/10/08 fix bug, 安静时长加入字典信息和列表
    def detect_gap(begin2, end1, speaker):
        '''
        当讲话者交换时，计算两次讲话之间的间隔时长，大于指定秒数才计算安静时间
        :param begin2: 本次讲话者开始的毫秒数
        :param end1: 上次讲话者结束的毫秒数
        :param speaker: 本次讲话者是谁： speaker1/speaker2
        修改的数据：  raw_dialog_list， 大于指定秒数的安静时间加入本次讲话者的讲话列表中
        :return:
        '''
        global file_content
        gap = round((begin2 - end1) / 1000) # 转换为秒数,四舍五入到个位
        # print('gap:', gap)
        if gap >= PAUSE_MIN_SECONDS:  # 安静N秒以上才计算安静时长
            file_content = ''.join([file_content, f'【安静{gap}秒】\n'])  # 字符串用join比+号效率高
            # 把当前讲话者讲的内容和时间戳加入对话列表
            raw_quiet_dict = {'begin': time_format(begin2), 'end': time_format(end1), 'speaker': speaker,
                              'text': f'【安静{gap}秒】'}
            raw_dialog_list.append(raw_quiet_dict)

    # 2020/10/08 adds:
    begin1 = begin2 = end1 = end2 = 0 #暂记录讲话者1和讲话者2的起止时间，用于计算两者讲话期间的安静时长

    while index < length:
        item = raw_list[index]
        # 讯飞转写好的格式：{"bg":"100","ed":"1220","onebest":"那你们怎么样？","speaker":"1"},
        # 记下每轮对话speaker1开始的时间戳

        if item['speaker'] == '1':
            begin1 = begin = int(item['bg'])

            detect_gap(begin, end, "speaker1") # 计算两个对话者中间的安静时长

            speaker1_quiet_time += (begin - end) / 1000

            item_s1['speaker'] = '1'
            item_s1['begin'] = time_format(int(item['bg']))  # 转换成 分:秒格式,例如：01:12
            str_s1 = ''
            # 循环读到speaker2为止
            while item['speaker'] == '1':
                str_s1 += item['onebest']  # 把同一个说话者一次说的话放在一起，直到另一说话者说话时为止
                item_s1['end'] = time_format(int(item['ed']))

                # 把当前讲话者讲的内容和时间戳加入列表
                '''
                raw_dialog_dict = {}  # 2020/07/02 Fix Bug.这里每次都得清空一次
                raw_dialog_dict['begin'] = time_format(int(item['bg']))
                raw_dialog_dict['end'] = time_format(int(item['ed']))
                raw_dialog_dict['speaker'] = 'speaker1'
                raw_dialog_dict['text'] = item['onebest']
                
                改写为以下形式：
                '''
                raw_dialog_dict = {'begin': time_format(int(item['bg'])), 'end': time_format(int(item['ed'])),
                                   'speaker': 'speaker1', 'text': item['onebest']}


                raw_dialog_list.append(raw_dialog_dict)
                # print(raw_dialog_list)
                end = int(item['ed'])
                t = (end - int(item['bg'])) / 1000  # 只计算说话秒数
                # print("speaker1's time:'", t)
                speaker1_time += t

                # 同一位讲话者，间隔一下之后继续说的话
                index += 1
                if index < length:
                    item = raw_list[index]  # 换下一段话
                    if item['speaker'] == '1':  # 同一讲话者
                        begin = int(item['bg'])
                        # print('speaker1:', time_format(int(item['bg'])))
                        # print('%d-%d=%d'%(begin, end, (begin-end)//1000))
                        # detect_gap(begin, end)    #用这个不能正常显示
                        gap = (begin - end) / 1000
                        # print('speaker1 gap:', gap)
                        speaker1_quiet_time += gap

                        end = int(item['ed'])  # new end                
                        gap = round(gap)
                        if gap >= PAUSE_MIN_SECONDS:  # 安静指定秒以上才显示
                            # str_s1 +=  '【安静%d秒】' %(gap) #不用安静起始的时间戳，文本看起来会清爽一些
                            str_s1 += item_s1['end'] + f'【安静{gap}秒】'  # 应该直接加到当前讲话者的内容里
                            # 2020/07/02 to add:
                            # append to list :安静时间
                            # 把当前讲话者安静的内容和时间戳加入列表

                            # 从上面改写如下（PyCharm建议）
                            raw_dialog_dict = {'begin': time_format(begin),
                                               'end': time_format(end),
                                               'speaker': 'speaker1',
                                               'text': '【安静%d秒】' % (gap)}
                            raw_dialog_list.append(raw_dialog_dict)
                    else:
                        break
                else:
                    break
            item_s1['text'] = str_s1
            # print("item_s1:", item_s1)

            # text = "%s  %s：%s\n" % (item_s1['begin'], speaker1, str_s1)
            # 2020/07/03 修改为起止时间戳：
            text = "%s-%s [%s]：%s\n" % (item_s1['begin'], item_s1['end'], speaker1, str_s1)

            # total_length_speaker1 += len(text)
            # total_length_speaker1 += count_zh_chars(text)
            # fix above bug, 多算了起止时间戳字符串
            total_length_speaker1 += count_zh_chars(str_s1)

            speaker1_content = ''.join([speaker1_content, text])
            file_content = ''.join([file_content, text])

            # end = int(item['ed'])
        # 记下每轮对话speaker2开始的时间戳
        if item['speaker'] == '2':
            begin = int(item['bg'])

            detect_gap(begin, end, "speaker2")  # 计算两个对话者中间的安静时长

            speaker2_quiet_time += (begin - end) / 1000

            item_s2['speaker'] = '2'
            item_s2['begin'] = time_format(int(item['bg']))
            str_s2 = ''
            # 循环读到speaker1为止
            while item['speaker'] == '2':
                str_s2 += item['onebest']
                item_s2['end'] = time_format(int(item['ed']))

                # 把当前讲话者讲的内容和时间戳加入列表
                raw_dialog_dict = {'begin': time_format(int(item['bg'])),
                                   'speaker': 'speaker2',
                                   'end': time_format(int(item['ed'])),
                                   'text': item['onebest']}
                raw_dialog_list.append(raw_dialog_dict)

                end = int(item['ed'])
                t = (end - int(item['bg'])) / 1000  # 只计算说话秒数
                # print("speaker2's time:'", t)
                speaker2_time += t

                index += 1
                if index < length:
                    item = raw_list[index]
                    if item['speaker'] == '2':
                        begin = int(item['bg'])
                        # detect_gap(begin, end)
                        gap = (begin - end) / 1000
                        # print('speaker2 gap:', gap)
                        speaker2_quiet_time += gap
                        end = int(item['ed'])  # new end
                        gap = round(gap)
                        if gap >= PAUSE_MIN_SECONDS:  # 安静指定秒以上才显示
                            str_s2 += item_s2['end'] + '【安静%d秒】' % (gap)  # 应该直接加到当前讲话者的内容里
                            # 2020/07/02 to add:
                            # append to list :安静时间
                            # 把当前讲话者安静的内容和时间戳加入列表
                            raw_dialog_dict = {'begin': time_format(begin),
                                               'end': time_format(end),
                                               'speaker': 'speaker2',
                                               'text': '【安静%d秒】' % (gap)}
                            raw_dialog_list.append(raw_dialog_dict)

                    else:
                        break
                else:
                    break
            item_s2['text'] = str_s2

            # text = "%s  %s：%s\n" % (item_s2['begin'], speaker2, str_s2)
            # 2020/07/03 修改为起止时间戳：
            text = "%s-%s [%s]：%s\n" % (item_s2['begin'], item_s2['end'], speaker2, str_s2)

            # total_length_speaker2 += len(text)
            # total_length_speaker2 += count_zh_chars(text)
            total_length_speaker2 += count_zh_chars(str_s2)

            speaker2_content = ''.join([speaker2_content, text])
            file_content = ''.join([file_content, text])

    total_time = time_format(int(item['ed']))

    # 调用函数，编辑讯飞语音中的常见错误：
    speaker1_content = edit_common_error(speaker1_content)
    speaker2_content = edit_common_error(speaker2_content)
    file_content = edit_common_error(file_content)

    # 内部函数：统计重要数据：说话时长、安静时长、字节数及相应占比
    def calc_statistics():

        # 计算两说话时长占比
        speaker1_time_ratio = speaker1_time / (speaker1_time + speaker2_time)
        # fw.write('speaker1:' + ' ' * 10 + '说话时长：%d秒,占比：%.2f%%  安静时长:%d 秒。\n' % (speaker1_time, speaker1_time_ratio * 100, output_result[3]))
        # fw.write('speaker2:' + ' ' * 10 + '说话时长：%d秒,占比：%.2f%%  安静时长:%d 秒。\n' % (speaker2_time, (1-speaker1_time_ratio) * 100, output_result[4]))

        # 计算两人说话字节占比
        speaker1_words_ratio = total_length_speaker1 / (total_length_speaker1 + total_length_speaker2)
        # fw.write('speaker1:' + ' ' * 10 + '字节总数：%d,占比：%.2f%%\n' % (total_length_speaker1, ratio * 100))
        # fw.write('speaker2:' + ' ' * 10 + '字节总数：%d,占比：%.2f%%\n' % (total_length_speaker2, (1 - ratio) * 100))

        speaker1_dict = OrderedDict()  # 需要用此，按顺序写入，以便之后按顺序读出
        speaker1_dict["说话者"] = "speaker1"
        speaker1_dict["说话时长"] = str(round(speaker1_time)) + "秒"
        speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
        speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
        speaker1_dict["字符总数"] = "%d" % total_length_speaker1
        speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)

        speaker2_dict = OrderedDict()
        speaker2_dict["说话者"] = "speaker2"
        speaker2_dict["说话时长"] = str(round(speaker2_time)) + "秒"
        speaker2_dict["时长占比"] = "%.2f%%" % ((1 - speaker1_time_ratio) * 100)
        speaker2_dict["安静时长"] = "%d 秒" % speaker2_quiet_time
        speaker2_dict["字符总数"] = "%d" % total_length_speaker2
        speaker2_dict["字符占比"] = "%.2f%%" % ((1 - speaker1_words_ratio) * 100)

        # print("speaker1:", speaker1_dict)
        # print("speaker2:", speaker2_dict)

        return speaker1_dict, speaker2_dict

    # print(raw_dialog_list)

    speaker1_stat_dict, speaker2_stat_dict = calc_statistics()

    result = (file_content, speaker1_content, speaker2_content,
              speaker1_stat_dict, speaker2_stat_dict, raw_dialog_list, total_time)

    return result


# raw,只按原始转录内容，加上时间戳
# 用途：方便用户对照着进行编辑，以纠正原始转录时可能出现的错误，例如：特别是搞错了是谁说的话
# 需要先调用convert_data函数,
# output_result是由函数convert_data返回的元组
# 默认是输出两位说话者的文字two， 可选择speaker1, or speaker2，只输出教练的文字
def write_raw2timestamp_txt(raw_txt_file, output_file):
    raw_list = get_result_list(raw_txt_file)  # 获取讯飞原始转写的数据信息

    try:
        fw = open(output_file, 'w', encoding='utf-8')
        all_text = ''
        for item in raw_list:
            begin = time_format(int(item['bg']))  # 按时间戳显示开始时间
            # 用Jet自定义函数处理一下原始转写的文字
            text = edit_common_error(item['onebest'])

            # text = item['onebest'] # 不处理，保留原始转写的结果
            # text = begin + '  speaker' + item['speaker'] + '：' + text + '\n'  # 不加换行，否则太多！
            text = begin + '  speaker' + item['speaker'] + '：' + text  # 不加换行，否则太多！

            # print(text)
            all_text = ''.join([all_text, text])
        fw.write(all_text)
        print('从原始文件转成时间戳文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


def get_stamp_list_from_rawtxt(raw_txt_file, coach='speaker1'):
    """
    Jet adds: 2020/07/20
    从原始txt文件，转换成带起止时间戳的list
    根据当前speaker1/2是教练/客户来改写

    version 2.0 格式：
    [{'speak_time': '00:01-00:03', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'speak_time': '00:03-00:05', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }


    :param raw_txt_file:
    :return: stamp_list
    """
    # 2020/07/03 修改用以下的数据，包含安静的提示
    output_result = convert_2speakers_rawtxt(raw_txt_file)  # 从原始转写文件获取信息到元组
    # 2020/07/25 为提升运行效率，只调用此函数一次，获取数据缓存
    #。。。

    raw_list = output_result[5]
    '''
    for i in range(len(raw_list)): #修改speaker1/speaker2 为 教练/客户
        if coach == "speaker1" and raw_list[i]['speaker'] =="speaker1":
             raw_list[i]['speaker'] = "教练"
        elif coach == "speaker2" and raw_list[i]['speaker'] =="speaker2":
             raw_list[i]['speaker'] = "教练"
        else:
             raw_list[i]['speaker'] = "客户"
    '''
    stampdict_list = []
    # 根据当前speaker1/2是教练/客户来改写
    for i, item in enumerate(raw_list):
        if coach == "speaker1" and item['speaker'] == "speaker1":
            item['speaker'] = "教练"
        elif coach == "speaker2" and item['speaker'] == "speaker2":
            item['speaker'] = "教练"
        else:
            item['speaker'] = "客户"

        start_time = item['begin']
        end_time = item['end']
        text = item['text']

        # 改时间戳格式为起止格式
        timestamp_dict = {
            "speak_time": start_time + "-" + end_time,  # begin_time 改为speak_time(起止时间戳）
            "speaker": item['speaker'],
            "content": text
        }
        stampdict_list.append(timestamp_dict)

    return output_result, stampdict_list


# 从原始转写txt，加上起止时间戳，写入二进制文件，保存对象。
def convert_raw2timestamp_dictlist_binary_file(raw_txt_file, output_file, coach='speaker1'):
    """
    Jet adds: 2020/06/30 以二进制方式写入文件，保存对象。测试成功
    把转写的原始txt文件，转换成带时间戳的dictlist文件

    2020/07/03
    添加参数coach

    2020/07/02
    注：忽略了安静的时间戳，需要加上结束的时间戳

    version 2.0 格式：
    [{'speak_time': '00:01-00:03', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'speak_time': '00:03-00:05', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }

    06/30
    version 1.0 格式：
    [{'start_time': '00:01', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'start_time': '00:03', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "start_time":"00:01:21",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }

    :param raw_txt_file:
    :param output_file:
    :return: dict_list
    """
    # raw_list = get_result_list(raw_txt_file)  # 获取讯飞原始转写的数据信息
    # 注： 以上函数没有修改原始数据！！
    # 2020/07/03 修改用以下的数据，包含安静的提示
    output_result = convert_2speakers_rawtxt(raw_txt_file)  # 从原始转写文件获取信息到元组
    raw_list = output_result[5]

    for i in range(len(raw_list)):  # 修改speaker1/speaker2 为 教练/客户
        if coach == "speaker1" and raw_list[i]['speaker'] == "speaker1":
            raw_list[i]['speaker'] = "教练"
        elif coach == "speaker2" and raw_list[i]['speaker'] == "speaker2":
            raw_list[i]['speaker'] = "教练"
        else:
            raw_list[i]['speaker'] = "客户"

    print("raw list:", raw_list)
    dict_list = []
    try:
        # fw = open(output_file, 'w', encoding='utf-8')
        fw = open(output_file, 'wb')
        for item in raw_list:
            # start_time = time_format(int(item['bg']))  # 按时间戳显示开始时间
            # end_time = time_format(int(item['ed']))  #添加截止时间戳

            start_time = item['begin']
            end_time = item['end']
            # 用Jet自定义函数处理一下原始转写的文字
            text = item['text']  # 先不用这个

            # 改时间戳格式为起止格式
            timestamp_dict = {
                "speak_time": start_time + "-" + end_time,  # begin_time 改为speak_time(起止时间戳）
                "speaker": item['speaker'],
                "content": text
            }
            dict_list.append(timestamp_dict)
        pickle.dump(dict_list, fw)
        # fw.write(str(dict_list))
        print('从原始文件转成时间戳dictlist二进制文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# Jet adds on 2020／06／30 , 测试成功
# 从二进制文件读取dict list对象
# 用途：可从返回的数据填充UI的表格，供用户查看，或编辑修改
def read_stamplist_from_binary_file(file_name):
    fr = open(file_name, 'rb')  # b代表用二进制方式读取文件
    dict_list = pickle.load(fr)

    return dict_list


# Jet adds on 2020／07／03 , 测试成功
# 从二进制文件读取icf dict list对象
# 用途：可从返回的数据填充UI的表格，供用户查看，或编辑修改
def read_icflist_from_binary(file_name):
    fr = open(file_name, 'rb')  # b代表用二进制方式读取文件
    dict_list = pickle.load(fr)
    return dict_list


# Jet adds: 2020/06/29
def convert_raw2timestamp_dictlist_txt_file(raw_txt_file, output_file):
    """
    把转写的原始txt文件，转换成带时间戳的dictlist文件

    07/02    version 2.0 格式：
    [{'speak_time': '00:01-00:03', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'speak_time': '00:03-00:05', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }


    06/29 v.1格式：
    [{'start_time': '00:01', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'start_time': '00:03', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "start_time":"00:01:21",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }

    :param raw_txt_file:
    :param output_file:
    :return: dict_list
    """
    raw_list = get_result_list(raw_txt_file)  # 获取讯飞原始转写的数据信息
    dict_list = []
    try:
        fw = open(output_file, 'w', encoding='utf-8')
        all_text = ''
        for item in raw_list:
            start_time = time_format(int(item['bg']))  # 按时间戳显示开始时间
            end_time = time_format(int(item['ed']))

            # 用Jet自定义函数处理一下原始转写的文字
            text = edit_common_error(item['onebest'])  # 先不用这个
            timestamp_dict = {
                "speak_time": start_time + "-" + end_time,
                "speaker": item['speaker'],
                "content": text
            }

            dict_list.append(timestamp_dict)

        fw.write(str(dict_list))
        print('从原始文件转成时间戳dictlist文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 只输出文本，不显示时间戳，只按1个讲话者
def write_raw2nonstamp_txt(raw_txt_file, output_file):
    raw_list = get_result_list(raw_txt_file)  # 获取讯飞原始转写的数据信息

    try:
        fw = open(output_file, 'w', encoding='utf-8')
        text = ''
        for item in raw_list:
            # begin = time_format(int(item['bg']))  # 按时间戳显示开始时间
            # 处理一下原始转写的文字
            text = ''.join([text, edit_common_error(item['onebest'])])
            # 不处理，保留原始转写的结果
            # print(text)

        fw.write(text)
        print('从原始文件转成时间戳文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 文件开头、统计字节数及笔录写入文件
# 需要先调用convert_data函数
# output_result是由函数convert_data返回的元组
# 默认是输出两位说话者的文字two， 可选择speaker1, or speaker2，只输出教练的文字
# def export_icf_with_stat_to_txt(output_file, output_result, speaker="two"):

# 2020/07/03 modify, need to debug result...
def export_icf_with_stat_to_txt(output_file, output_result, coach="speaker1"):
    try:
        fw = open(output_file, 'w', encoding='utf-8')
        # 写文件开头
        text = '教练对话语音转写\n\n'
        # text += '注：请根据实际情况，替换speaker1或speaker2为教练或客户：\n'
        audio_time = output_result[6]

        text += '总时长：%s\n\n' % audio_time  # total_time is calculated in function: convert_data

        # 打印标题统计speaker1, speaker2的讲话及安静信息：
        # text = ''
        '''
        #参考统计信息的字典格式：
        speaker1_dict = OrderedDict()  # 需要用此，按顺序写入，以便之后按顺序读出
        speaker1_dict["说话者"] = "speaker1"
        speaker1_dict["说话时长"] = str(round(speaker1_time)) + "秒"
        speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
        speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
        speaker1_dict["字符总数"] = "%d" % total_length_speaker1
        speaker1_dict["字符占比"] = "%.2f%%\n" % (speaker1_words_ratio * 100)
        '''
        # 根据传入的实际参数，转换speaker1/speaker2 为"教练/客户"
        speaker1_dict = output_result[3]
        length = 0
        for key in speaker1_dict:
            value = speaker1_dict[key]
            if key == "说话者":
                if coach == "speaker1" and value == "speaker1":
                    value = "教练"
                elif coach == "speaker2" and value == "speakers":
                    value = "教练"
                else:
                    value = "客户"
            text = ''.join([text, key, value])
            length += 1
            if length < len(speaker1_dict):
                text += '，'

        text += '\n'

        speaker2_dict = output_result[4]
        length = 0
        for key in speaker2_dict:
            value = speaker2_dict[key]
            if key == "说话者":
                if coach == "speaker1" and value == "speaker1":
                    value = "教练"
                elif coach == "speaker2" and value == "speakers":
                    value = "教练"
                else:
                    value = "客户"
            text = ''.join([text, key, value])
            length += 1
            if length < len(speaker2_dict):
                text += '，'

        text += '\n'

        fw.write(text)  # 写标题和统计信息

        # 写ICF格式文字， 参数来自convert_data函数返回值的第一个
        fw.write(output_result[0])
        print("保存结果到{}成功。".format(output_file))
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 根据原始转写文件，转换成ICF格式，同一个人说话的内容合并到一个时间段显示
# 返回列表，包含讲话者安静的时间戳
def convert_rawfile2icf_format_list(raw_txt_file, coach="speaker1"):
    """
    根据参数，把speaker1/speaker2 转换为 教练/客户
    调用此函数的函数中，根据用户选择，确定哪位是教练
    参见convert_data函数中的定义:
    raw_dialog_dict['begin'] = time_format(int(item['bg']))
    raw_dialog_dict['end'] = time_format(int(item['ed']))
    raw_dialog_dict['speaker'] = 'speaker1'
    raw_dialog_dict['text'] = item['onebest']
    raw_dialog_list.append(raw_dialog_dict)

    :returns:
    icf_list:

    icf_stat_list:
    speaker1_dict["说话者"] = "教练"或"客户"
    speaker1_dict["说话时长"] = "t秒"
    speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
    speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
    speaker1_dict["字符总数"] = "%d" % total_length_speaker1
    speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)

    """

    output_result = convert_2speakers_rawtxt(raw_txt_file)  # 从原始转写文件获取信息到元组
    raw_list = output_result[5]

    for i in range(len(raw_list)):  # 修改speaker1/speaker2 为 教练/客户
        if coach == "speaker1" and raw_list[i]['speaker'] == "speaker1":
            raw_list[i]['speaker'] = "教练"
        elif coach == "speaker2" and raw_list[i]['speaker'] == "speaker2":
            raw_list[i]['speaker'] = "教练"
        else:
            raw_list[i]['speaker'] = "客户"

    icf_list = []  # ICF格式，对话List
    # 修改起止时间戳，以便显示到table widget中
    # 合并同一位讲话者的内容到同一时间戳内：
    i = 0
    coaching_row_count = 0  # 计算合并后的教练记录条数（行数）
    while i < len(raw_list):
        start_time = raw_list[i]['begin']  # 按时间戳显示开始时间

        speaker = raw_list[i]['speaker']
        if speaker == "教练":
            coaching_row_count += 1

        text = raw_list[i]['text']
        if (i + 1) < len(raw_list):  # 避免list index out of range
            next_speaker = raw_list[i + 1]['speaker']

            # 循环合并同一位讲话者的内容
            # 2020-07-20 fix bug: list index out or range
            # while speaker == next_speaker and (i + 1) < len(raw_list):
            while speaker == next_speaker and (i + 2) < len(raw_list):
                i += 1
                text += raw_list[i]['text']
                speaker = raw_list[i]['speaker']
                next_speaker = raw_list[i + 1]['speaker']

        end_time = raw_list[i]['end']  # 添加截止时间戳

        timestamp_dict = {
            "speak_time": start_time + "-" + end_time,  # begin_time 改为speak_time(起止时间戳）
            "speaker": speaker,
            "content": text
        }
        icf_list.append(timestamp_dict)
        i += 1

    icf_stat_list = []  # ICF教练/客户数据统计

    '''
    #参考统计信息的字典格式 in convert_data：
    speaker1_dict = OrderedDict()  # 需要用此，按顺序写入，以便之后按顺序读出
    speaker1_dict["说话者"] = "speaker1"  #需要转换成：教练或客户
    speaker1_dict["说话时长"] = str(round(speaker1_time)) + "秒"
    speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
    speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
    speaker1_dict["字符总数"] = "%d" % total_length_speaker1
    speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)
    '''
    # 根据传入的实际参数，转换speaker1/speaker2 为"教练/客户"
    # to do: modify, 在这里重新用icf_list计算说话合并后的字符总数，不再使用原始的函数

    speaker1_dict = output_result[3]
    if coach == "speaker1":
        speaker1_dict["说话者"] = "教练"
    else:
        speaker1_dict["说话者"] = "客户"

    speaker2_dict = output_result[4]
    if coach == "speaker2":
        speaker2_dict["说话者"] = "教练"
    else:
        speaker2_dict["说话者"] = "客户"

    icf_stat_list.append(speaker1_dict)
    icf_stat_list.append(speaker2_dict)

    # 查看是否修改的教练/客户是否正确。 OK，成功！
    # for item in icf_list:
    #    print(item)

    return icf_list, icf_stat_list, coaching_row_count

# 为了避免重复调用，修改以上函数的参数
# 根据原始转写文件，转换成ICF格式，同一个人说话的内容合并到一个时间段显示
# 返回列表，包含讲话者安静的时间戳
# output_result = convert_2speakers_rawtxt(raw_txt_file)  # 从原始转写文件获取信息到元组
#
def convert_raw2icf_format_list(output_result, coach="speaker1"):
    """
    调用此函数的函数中，根据用户选择，确定哪位是教练
    参见convert_data函数中的定义:
    raw_dialog_dict['begin'] = time_format(int(item['bg']))
    raw_dialog_dict['end'] = time_format(int(item['ed']))

    2020/07/25
    特别注意：与convert_rawfile2icf_format_list(raw_txt_file, coach="speaker1")不同的是：
    函数 transform_tool.get_stamp_list_from_rawtxt
    已经把speaker1/speaker2 转换为 教练/客户

    raw_dialog_dict['speaker'] = '教练/客户'

    raw_dialog_dict['text'] = item['onebest']
    raw_dialog_list.append(raw_dialog_dict)

    :returns:
    icf_list:

    icf_stat_list:
    speaker1_dict["说话者"] = "教练"或"客户"
    speaker1_dict["说话时长"] = "t秒"
    speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
    speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
    speaker1_dict["字符总数"] = "%d" % total_length_speaker1
    speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)

    """
    #获取缓存中从raw.txt获取的raw_dialog_list
    raw_list = output_result[5]

    '''
    2020/07/25
    特别注意：与convert_rawfile2icf_format_list(raw_txt_file, coach="speaker1")不同的是：
    函数 transform_tool.get_stamp_list_from_rawtxt
    已经把speaker1/speaker2 转换为 教练/客户
    for i in range(len(raw_list)):  # 修改speaker1/speaker2 为 教练/客户
        if coach == "speaker1" and raw_list[i]['speaker'] == "speaker1":
            raw_list[i]['speaker'] = "教练"
        elif coach == "speaker2" and raw_list[i]['speaker'] == "speaker2":
            raw_list[i]['speaker'] = "教练"
        else:
            raw_list[i]['speaker'] = "客户"
    '''

    icf_list = []  # ICF格式，对话List
    # 修改起止时间戳，以便显示到table widget中
    # 合并同一位讲话者的内容到同一时间戳内：
    i = 0
    coaching_row_count = 0  # 计算合并后的教练记录条数（行数）
    while i < len(raw_list):
        start_time = raw_list[i]['begin']  # 按时间戳显示开始时间

        speaker = raw_list[i]['speaker']
        if speaker == "教练":
            coaching_row_count += 1

        text = raw_list[i]['text']
        if (i + 1) < len(raw_list):  # 避免list index out of range
            next_speaker = raw_list[i + 1]['speaker']

            # 循环合并同一位讲话者的内容
            # 2020-07-20 fix bug: list index out or range
            # while speaker == next_speaker and (i + 1) < len(raw_list):
            while speaker == next_speaker and (i + 2) < len(raw_list):
                i += 1
                text += raw_list[i]['text']
                speaker = raw_list[i]['speaker']
                next_speaker = raw_list[i + 1]['speaker']

        end_time = raw_list[i]['end']  # 添加截止时间戳

        timestamp_dict = {
            "speak_time": start_time + "-" + end_time,  # begin_time 改为speak_time(起止时间戳）
            "speaker": speaker,
            "content": text
        }
        icf_list.append(timestamp_dict)
        i += 1

    icf_stat_list = []  # ICF教练/客户数据统计

    '''
    #参考统计信息的字典格式 in convert_data：
    speaker1_dict = OrderedDict()  # 需要用此，按顺序写入，以便之后按顺序读出
    speaker1_dict["说话者"] = "speaker1"  #需要转换成：教练或客户
    speaker1_dict["说话时长"] = str(round(speaker1_time)) + "秒"
    speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
    speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
    speaker1_dict["字符总数"] = "%d" % total_length_speaker1
    speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)
    '''
    # 根据传入的实际参数，转换speaker1/speaker2 为"教练/客户"
    # to do: modify, 在这里重新用icf_list计算说话合并后的字符总数，不再使用原始的函数

    speaker1_dict = output_result[3]
    if coach == "speaker1":
        speaker1_dict["说话者"] = "教练"
    else:
        speaker1_dict["说话者"] = "客户"

    speaker2_dict = output_result[4]
    if coach == "speaker2":
        speaker2_dict["说话者"] = "教练"
    else:
        speaker2_dict["说话者"] = "客户"

    icf_stat_list.append(speaker1_dict)
    icf_stat_list.append(speaker2_dict)

    # 查看是否修改的教练/客户是否正确。 OK，成功！
    # for item in icf_list:
    #    print(item)

    return icf_list, icf_stat_list, coaching_row_count


# 把以上函数获得的icf_list存到二进制文件中
#def save_raw2icf_list_binary(raw_txt_file, output_file, coach="speaker1"):
#icf_list = convert_raw2icf_format_list(raw_txt_file, coach)
# 2020-/07/25 modified
def save_raw2icf_list_binary(icf_list, output_file, coach="speaker1"):
    """
    Jet adds: 2020/06/30 以二进制方式写入文件，保存对象。测试成功
    把转写的原始txt文件，转换成带时间戳的dictlist文件

    2020/07/02
    注：忽略了安静的时间戳，需要加上结束的时间戳

    version 2.0 格式：
    [{'speak_time': '00:01-00:03', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'speak_time': '00:03-00:05', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }

    :param raw_txt_file:
    :param output_file:
    :return: dict_list
    """
    try:
        fw = open(output_file, 'wb')  # 二进制文件，用于写入list对象
        pickle.dump(icf_list, fw)

        print('从原始文件转成icf时间戳dictlist二进制文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()

    return icf_list


# 需要先调用convert_data函数
# 自动生成word文档，包括:标题、统计信息、对话时间戳及对话内容（教练的话以蓝色字体标记）
def write_to_word_file(output_file, output_result, speaker="two", coach="speaker1"):
    # 文件名称
    file_name = output_file
    # 打开文件
    document = docx.Document()

    # 添加指定格式的段落
    # text：段落内容
    def add_formatted_text(document, text, font_cn_name="宋体", size=12, color=RGBColor(0, 0, 0), alignment="left",
                           is_bold=False):
        # 添加段落
        p1 = document.add_paragraph()
        text1 = p1.add_run(text)

        text1.font.name = 'Times New Roman'  # 控制是西文时的字体
        text1.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn_name)  # 控制是中文时的字体

        text1.font.size = Pt(size)
        text1.font.color.rgb = color

        # 此处高亮显示“安静”为绿色， 如果有的话
        # 此处还需改进。。。教练的文字，字体不对，颜色蓝色也没有了
        # highlight_keywords(p1, text, keyword, color)

        # 加粗
        text1.font.bold = is_bold
        # 斜体
        # text1.font.italic = True

        if alignment == "center":
            p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == "left":
            p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p1.paragraph_format.space_after = Pt(10)  # 段后间距，default

        document.save(file_name)

    # 第一段（设置标题居中）
    title = '教练对话语音转写'
    add_formatted_text(document, title, font_cn_name="宋体", size=16, color=RGBColor(0, 0, 255), alignment="center",
                       is_bold=True)

    # remind ='注：请根据实际情况，替换speaker1或speaker2为教练或客户。'
    # add_formatted_text(document, remind)

    total_time = '总时长：59分钟'
    add_formatted_text(document, total_time)

    # 添加表格, 显示统计教练和客户的信息
    table = document.add_table(rows=3, cols=6,
                               style="Light List Accent 5")  # 指定表格格式，参见：https://blog.csdn.net/ibiao/article/details/78595295
    table.style.font.name = '宋体'  # 表格内的字体类型
    table.style.font.size = 120000  # 字号

    # 在表的指定行，列，输入内容
    def add_text2table(table, row, col, text):
        cell = table.cell(row, col)
        cell.text = text

    # 需要用到排序字典，#循环遍历，添加入表格：
    speaker1_dict = output_result[3]
    speaker2_dict = output_result[4]

    row, col = 0, 0
    # 表格各列标题
    for key in speaker1_dict:
        add_text2table(table, row, col, key)
        col += 1

    row, col = 1, 0
    for key in speaker1_dict:
        if col == 0:  # 根据参数决定哪位是教练，默认第一位是教练
            if coach == "speaker1":
                value = "教练"
            else:
                value = "客户"
        else:
            value = speaker1_dict[key]
            # print(key, value)
        add_text2table(table, row, col, value)
        col += 1

    row, col = 2, 0
    for key in speaker2_dict:
        if col == 0:  # 根据参数决定哪位是教练，默认第一位是教练
            if coach == "speaker2":
                value = "教练"
            else:
                value = "客户"
        else:
            value = speaker2_dict[key]
            # print(key, value)
        add_text2table(table, row, col, value)
        col += 1

    text = "笔录:"
    add_formatted_text(document, text)

    # 获取原始数据，转换成列表，以便操作
    speaking_list = output_result[0].split('\n')
    for item in speaking_list:
        if "speaker1" in item:
            if speaker == "speaker2":  # 跳过，不输出
                continue
            if coach == "speaker1":
                # 教练文字用蓝色、加粗
                item = item.replace("speaker1", "教练")

                add_formatted_text(document, item, color=RGBColor(0, 0, 255), is_bold=True)

                # print("coach:", item)
            else:
                item = item.replace("speaker1", "客户")
                add_formatted_text(document, item)
        elif "speaker2" in item:
            if speaker == "speaker1":  # 跳过，不输出
                continue
            if coach == "speaker2":
                # 教练文字用蓝色、加粗
                item = item.replace("speaker2", "教练")
                add_formatted_text(document, item, color=RGBColor(0, 0, 255), is_bold=True)
                # print("coach:", item)
            else:
                item = item.replace("speaker2", "客户")
                add_formatted_text(document, item)

    print("保存结果到{}成功。".format(output_file))


# 按照讯飞的转写结果显示，同一人的讲话有停顿会分开，（后续可以让用户手动编辑，删除多余的时间戳）
# 避免因讯飞识别发音人出错而导致时间戳错误
def get_result_list(filename):
    try:
        fr = open(filename, 'r', encoding='utf-8')

        str_data = fr.read()
        data_dict = eval(str_data)  # 字符串转换为字典类型
        # print(data_dict)
        result_list = json.loads(data_dict['data'])  # 将json格式的字符串转为列表
        # print(result_list)
        print("从文件{}读取数据成功。".format(filename))
        return result_list
    except:
        print("从文件{}读取数据失败！".format(filename))


# 把按轮流对话的方式存入文件，字典的列表（字符串）格式
def save_to_file(filename):
    try:
        fw = open(filename, 'w', encoding='utf-8')
        fw.write(str(convert_2speakers_rawtxt(filename)))
        print("保存结果到{}成功。".format(filename))
    except:
        print("保存失败!")
    finally:
        fw.close()


# 从指定的icf_list, icf_stat_list, 转写到指定的word文档
def save_icf_list_to_word(icf_list, icf_stat_list, file_name):
    # ----------------
    try:
        # 打开文件
        document = docx.Document()

        # 以下会导致程序崩溃！
        # document = docx.Document("ICF-Word-template.docx")  #用Jet自定义的模板（带页眉和页脚，有页码）
        # '''添加页码，以下代码OK，注意需要import OxmlElement和os

        # 在页脚添加页码
        jets_docx.add_page_number(document.sections[0].footer.paragraphs[0])
        document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.save(file_name)
        # '''

        # 第一段（设置标题居中）
        title = '教练对话语音转写'
        jets_docx.add_formatted_text(document, file_name, title, font_cn_name="宋体", size=16,
                                     color=RGBColor(0, 0, 255), alignment="center",
                                     is_bold=True)

        str_time = icf_list[len(icf_list) - 1]["speak_time"]  # 最后一个记录的起止时间，eg: 59:37-59:39
        audio_time = str_time.split('-')[1]
        total_time = '教练时长：%s' % audio_time
        jets_docx.add_formatted_text(document, file_name, total_time)

        # 是否输出统计信息：
        if icf_stat_list:  # 确保已经有统计数据
            # if icf_stat_list is not None:  # 确保已经有统计数据

            # 添加表格, 显示统计教练和客户的信息
            table = document.add_table(rows=3, cols=6,
                                       style="Light List Accent 5")
            # 指定表格格式，参见：https://blog.csdn.net/ibiao/article/details/78595295

            table.style.font.name = '宋体'  # 表格内的字体类型
            table.style.font.size = 120000  # 字号

            # 在表的指定行，列，输入内容
            def add_text2table(table, row, col, text):
                cell = table.cell(row, col)
                cell.text = text

            # 表格各列标题
            row, col = 0, 0
            for key in icf_stat_list[0]:
                add_text2table(table, row, col, key)
                col += 1

            # 分别打印教练和客户的统计信息
            row, col = 1, 0
            for value in icf_stat_list[0].values():  # 注意：必须加上.values(),否则只返回keys
                add_text2table(table, row, col, value)
                col += 1

            row, col = 2, 0
            for value in icf_stat_list[1].values():
                add_text2table(table, row, col, value)
                col += 1

        text = "教练对话笔录:"
        jets_docx.add_formatted_text(document, file_name, text)

        # 用户在右侧icf_table中看到什么，就打印什么（只有教练的讲话或包括教练和客户的对话）
        for item_dict in icf_list:  # 分别打印教练和客户的统计信息
            '''
                设置格式为：
                00:01-00:02[教练] 喂，谭*你好！
                00:03-00:05[客户] Hello。李**
                '''
            text = ''.join([item_dict["speak_time"], ' [', item_dict["speaker"], '] ', item_dict["content"]])
            # 把教练讲话的字体设为蓝色：
            if item_dict["speaker"] == "教练":
                jets_docx.add_formatted_text(document, file_name, text, color=RGBColor(0, 0, 255))
            # elif item_dict[
            #    "speaker"] == "客户" and not self.ckbox_only_coach_text:  # 客户讲话的字体用黑色,如果选中只输出教练文字时，则不输出客户文字
            # 07/09 fix above bug by adding .isChecked()

            # 客户讲话的字体用黑色,如果选中只输出教练文字时，则不输出客户文字
            elif item_dict["speaker"] == "客户":
                jets_docx.add_formatted_text(document, file_name, text)

        print("保存结果到{}成功。".format(file_name))
        return True
    except IOError:
        print("保存结果到{}失败。".format(file_name))
        return False


# 无需进度条，命令行调用，无GUI
# 1个讲话者,可用于视频转写字幕（还需加入时间戳）， 讲章预备, 讲章回顾等
def convert_1speaker_audio_to_text(audio_file, output_file):
    """
    :param audio_file:
    :param output_file:纯文本，无时间戳
    :return:
    """
    start = datetime.datetime.now()
    print("开始转写" + start.strftime('%Y-%m-%d_%H_%M_%S'))

    api = XunfeiAudioConvertRequestApi(audio_file, speaker_number=1)
    result = api.all_api_request()

    # 执行批量修改
    edited_text = edit_common_error(str(result))
    # 指定编码格式为UTF-8格式，中英文都可正常存储
    try:
        fw = open(output_file, 'w', encoding='utf-8')  # w为写文件write
        fw.write(edited_text)
        msg = "语音转写内容保存到文件:%s成功！" % output_file
    except:
        msg = "语音转写内容保存到文件:%s失败" % output_file
    finally:
        fw.close()

    end = datetime.datetime.now()
    print("结束转写" + end.strftime('%Y-%m-%d_%H_%M_%S'))
    minutes, seconds = divmod((end - start).seconds, 60)
    print(f"总共用时：{minutes}分{seconds}秒。")

# 转写1人讲话原始文件为纯文本
def get_1speaker_txtonly(raw_file):
    """
    :param raw_file: 已经转写好的原始文件 
    :return: 
    """
    start = datetime.datetime.now()
    print("开始转写" + start.strftime('%Y-%m-%d_%H_%M_%S'))

    text = convert_1speaker_timestamp_rawtxt(raw_file)[1]

    text_file = raw_file.rpartition('.')[0] + '_textonly.txt'

    with open(text_file, 'w', encoding='utf-8') as f:
        f.write(text)

    end = datetime.datetime.now()
    print("结束转写" + end.strftime('%Y-%m-%d_%H_%M_%S'))
    minutes, seconds = divmod((end - start).seconds, 60)
    print(f"总共用时：{minutes}分{seconds}秒。")

# 转写1人讲话原始文件为纯文本,说话有间隔的，就添加换行分段
def get_1speaker_txt_passages(raw_file):
    """
    :param raw_file: 已经转写好的原始文件
    :return:
    """
    start = datetime.datetime.now()
    print("开始转写" + start.strftime('%Y-%m-%d_%H_%M_%S'))

    text = ''
    text_list = convert_1speaker_timestamp_rawtxt(raw_file)[2]
    for item in text_list:
        # 参见 convert_1speaker_timestamp_rawtxt函数中的定义
        text = ''.join([text, item['text'], '\n'])  # 添加分段的换行

    text_file = raw_file.rpartition('.')[0] + '_text_passages.txt'

    with open(text_file, 'w', encoding='utf-8') as f:
        f.write(text)

    end = datetime.datetime.now()
    print("结束转写" + end.strftime('%Y-%m-%d_%H_%M_%S'))
    minutes, seconds = divmod((end - start).seconds, 60)
    print(f"总共用时：{minutes}分{seconds}秒。")


# Jet adds:
def main():
    # audio (mp3) to raw.txt
    # 转写录制的讲章大纲
    audio_file = r"C:\#1-Jet's Docs\2020-Ministry\James\Jam1.12-18-Outline.mp3"
    raw_file = audio_file.rpartition('.')[0] + '.txt'
    # convert_1speaker_audio_to_text(audio_file, raw_file) # 提交获得原始转写文字
    get_1speaker_txt_passages(raw_file)

    '''
    output_file = raw_txt_file + "raw_timestamp.txt"
    save_1speaker_raw_timestamp2txt(raw_txt_file, output_file)

    output_file = raw_txt_file + "combine_timestamp.txt"
    save_1speaker_combined_timestamp2txt(raw_txt_file, output_file)

    output_file = raw_txt_file + "only_text.txt"
    save_1speaker_only_txt(raw_txt_file, output_file)
    '''
    # output_file = raw_txt_file + "raw.txt"
    '''
        version 2.0 格式：
    [{'speak_time': '00:01-00:03', 'speaker': '1', 'content': '喂，谭*你好！'},
    {'speak_time': '00:03-00:05', 'speaker': '2', 'content': ' Hello。'},]

    timestamp_dict = {
        "speak_time":"01:21-01:25",
        "speaker": "speaker1",
        "content": "what the speaker says"
    }
    '''


    """
    # 测试转写成timestamp格式纯文本文件
    raw_file = r"F:\Jet Python Study\PyCharmProjects\icf-audio2text\src\2010-06-02-coachingTanche_Raw.txt"
    output_file = raw_file + "2timestamp_binary.data"
    convert_raw2timestamp_dictlist_binary_file(raw_file, output_file)  # 写入

    #output_file = r"F:\Jet Python Study\PyCharmProjects\icf-audio2text\src\2010-06-02-coachingTanche_Raw.txt2timestamp_binary.data"
    read_stamplist_from_binary_file(output_file)  # 读取
    """
    '''
    raw_txt_file = r"F:\Jet Python Study\PyCharmProjects\icf-audio2text\src\2010-06-02-coachingTanche_Raw.txt"
    convert_raw2icf_format_list(raw_txt_file, coach="speaker1")

    output_result = convert_2speakers_rawtxt(raw_txt_file) #从原始转写文件获取信息到元组
    #输出成文本文件

    output_file = raw_txt_file + '-整理_教练和客户.txt'
    '''

    '''

    output_file = raw_txt_file + '-整理_教练only.txt'
    write_to_txt_file(output_file, output_result, speaker='speaker1')
    
    output_file = raw_txt_file + '-整理_客户only.txt'
    write_to_txt_file(output_file, output_result, speaker='speaker2')
    
    #输出成word文件
    output_word = raw_txt_file + '-整理_教练和客户.doc'
    write_to_word_file(output_word, output_result, speaker='two', coach='speaker1')
    
    output_word = raw_txt_file + '-整理_教练only.doc'
    write_to_word_file(output_word, output_result, speaker='speaker1', coach='speaker1')
    
    '''

    '''
    output_word = raw_txt_file + '-整理_客户only.doc'
    write_to_word_file(output_word, output_result, speaker='speaker2', coach='speaker2')
    '''

    # save_to_file(output_file)


if __name__ == '__main__':
    main()
