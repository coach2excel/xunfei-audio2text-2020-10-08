# 从讯飞转写好的原始数据文件读取信息，转换成教练、客户对话的格式
# 包括时间戳
# 再用word模板，生成word文件

import json  # json字符串转换为列表类型时要用到
import datetime  # 没用到
import re  # regular expression, 字符串替换时要用到
from collections import OrderedDict  # 导入collections.OrderedDict

# 以下生成word文档需要
'''
Jet: 不能安装pip install docx,否则运行时会提示错误：
  File "C:\lib\site-packages\docx.py", line 30, in <module>
    from exceptions import PendingDeprecationWarning
ModuleNotFoundError: No module named 'exceptions'
If already installed docx, need to uninstall it: 
pip unistall docx

Then install python-docx on PyCharm:
pip install python-docx

'''
import docx
from docx.shared import Inches, Cm, Pt
from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 字体设置

# 修改word关键词颜色
import os
from docx.oxml.ns import qn  # 设置中文字体

# 统计字符数
import string
from collections import namedtuple

# audio_file = "./audio/2020-06-02-Sherry-facilitator.mp3"


# "JetCoachDeng-60m.txt"
# data_file = '../output/' +"2010-06-02-coach 谭澈zoom.wav-讯飞转写原始结果2.txt"

# "2020-06-01-coach礼燕-zhumu.wav-讯飞转写原始结果.txt"  # 讯飞语音已经转写好的文件

speaker1 = 'speaker1'  # 可GUI获取，修改值为教练/客户
speaker2 = 'speaker2'  # 可GUI获取，修改值为教练/客户
PAUSE_MIN_SECONDS = 3  # 安静3秒以上才显示
total_time = 0  # 音频总时长
total_length_speaker1 = 0
total_length_speaker2 = 0
file_content = ''  # 返回的全部对话文本内容（包括教练和客户），之后写入指定文件
speaker1_content = ''  # 仅保存speaker1的对话
speaker2_content = ''  # 仅保存speaker2的对话
speaker1_time = 0  # 讲话时间秒数
speaker2_time = 0  # 讲话时间秒数
speaker1_quiet_time = 0  # 安静时间秒数
speaker2_quiet_time = 0  # 安静时间秒数
keyword = "安静"  # 用于高亮显示的关键词

raw_dialog_list = []  # 带有时间戳的对话列表，以便后续对照编辑


# 转换成 分:秒格式,例如：01:12
def time_format(miniseconds):
    seconds = miniseconds / 1000
    # result = datetime.timedelta(seconds)#只保留到秒数
    m, s = divmod(seconds, 60)  # 利用元组进行整除、余数计算
    # h, m = divmod(m, 60) #小时，可根据需要添加
    # result = "%02d:%02d:%02d" % (h, m, s)
    result = "%02d:%02d" % (m, s)
    # print("time:", result)
    return result


'''
# 用正则表达式替换字典中指定的字符串 key-〉value
re.sub(`pattern`, `repl`, `string`, `count=0`, `flags=0`)

`pattern`, `repl`, `string` 为必选参数
`count`, `flags` 为可选参数
`pattern`正则表达式
`repl`被替换的内容，可以是字符串，也可以是函数
`string`正则表达式匹配的内容
`count`由于正则表达式匹配的结果是多个，使用count来限定替换的个数从左向右，默认值是0，替换所有的匹配到的结果
`flags`是匹配模式，`re.I`忽略大小写，`re.L`表示特殊字符集\w,\W,\b,\B,\s,\S，`re.M`表示多行模式，`re.S` ‘.’包括换行符在内的任意字符，`re.U`表示特殊字符集\w,\W,\b,\B,\d,\D,\s,\D
'''


def multiple_replace(text, adict):
    rx = re.compile('|'.join(map(re.escape, adict)))

    def one_xlat(match):
        return adict[match.group(0)]

    return rx.sub(one_xlat, text)


# 替换讯飞语音转写的常见错误
def edit_common_error(text):
    # 此处需要根据常见问题，添加到字典里...
    error_dict = {
        "。呢": "呢？",
        "？呢": "呢？",
        "，呢": "呢，",
        "？啊": "啊？",
        "。啊": "啊。",
        "，啊": "啊，",
        "，吧": "吧，",
        "？呀": "呀？",
        "。吧": "吧。",
        "好，呀": "好呀，",
        "，哈": "哈，",
        "。哈": "哈。",
        "很好。": "很好!",
        "非常好。": "非常好!",
        "。唉": "唉。",
        # "嗯":"嗯，",
        "，嘛": "嘛，",

    }
    if text == "嗯":  # 一句话只有这一个字
        text = "嗯。"
    elif "嗯" in text and len(text) > 1:
        text = "嗯，"
    else:
        text = multiple_replace(text, error_dict)
    return text


# 统计字符数（只包括中文, 不包括英文字母、数字，由于是说话，不包括标点符号）
# Jet：这比用len计算要准确得多！且教练说话占比会明显减少。
# 每小时计数要减少3000字！！！
# Word计算时把中文标点当作汉字了。
def count_chars(s):
    count_en = count_dg = count_sp = count_zh = count_pu = 0
    s_len = len(s)
    for c in s:
        if c in string.ascii_letters:
            count_en += 1
        elif c.isdigit():
            count_dg += 1
        elif c.isspace():
            count_sp += 1
        elif c.isalpha():
            count_zh += 1
        else:
            count_pu += 1
    # total_chars = count_zh + count_en + count_sp + count_dg + count_pu
    # total_chars = count_zh + count_en +  count_dg
    total_chars = count_zh

    return total_chars


# 设置Word（doc）文档中指定关键词的颜色，高亮显示：
def highlight_keywords(paragraph, content, keyword, color=RGBColor(0, 0, 0)):
    if content.find(keyword) != -1:  # 如果该行有关键字,就以关键字为分界进行分割
        pt = keyword
        res = re.split(pt, content)  # 分割

        # add_run在同一段添加内容
        run = paragraph.add_run(res[0])  # 输入关键字之前的字符
        run.font.color.rgb = color  # 字体颜色
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # print(res[0],end=' ')

        run = paragraph.add_run(keyword)  # 输入关键字
        run.font.name = u'黑体'  # 设置插入的字体
        run.font.size = Pt(12)
        # 设置关键字之后也的字体颜色，这里设置为绿色
        run.font.color.rgb = RGBColor(0, 250, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # print(res[1],end=' ')
        run = paragraph.add_run(res[1])  # 输入关键字之后的字符
        run.font.color.rgb = color  # 字体颜色
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(12)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    else:
        run = paragraph.add_run(content)  # 如果该行没有关键字，则直接输入word
        # print(content)
        run.font.name = u'宋体'  # 设置插入的字体
        run.font.size = Pt(15)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


# 转换成列表（需要按对话顺序）
# 同一个人的讲话，合并到一起，直到换另一个讲话
# 注：如果转写结果本身把说话者搞错了，那么结果就错得厉害，并且连时间戳也会出问题！
# 二人对话的中间，可能有沉默时间！
def convert_data(data_file):
    global file_content
    global total_time
    global total_length_speaker1
    global total_length_speaker2
    global speaker1_content
    global speaker2_content
    global speaker1_time
    global speaker2_time
    global speaker1_quiet_time
    global speaker2_quiet_time
    global raw_dialog_list

    # 计算两次讲话之间的间隔时长>3秒才算
    def detect_gap(begin, end):
        global file_content
        gap = (begin - end) / 1000
        # print('gap:', gap)
        if round(gap) >= PAUSE_MIN_SECONDS:  # 安静3秒以上才显示
            file_content = ''.join([file_content, '【安静%d秒】\n' % (gap)])  # 字符串用join比+号效率高

    raw_list = get_result_list(data_file)  # 获取讯飞原始转写的数据信息
    # print('raw:', raw_list)
    conversation_list = []  # 记录所有对话，按speaker1, speaker2交替的形式存储, 没用上！
    index = 0
    length = len(raw_list)
    item_s1 = {}  # 记录每次对话speaker1说的时间戳和内容
    item_s2 = {}  # 记录每次对话speaker2说的时间戳和内容
    end = 0

    raw_dialog_dict = {}  # OrderedDict()#后续访问需要按添加的顺序访问
    raw_dialog_list = []

    while index < length:
        item = raw_list[index]
        # 讯飞转写好的格式：{"bg":"100","ed":"1220","onebest":"那你们怎么样？","speaker":"1"},
        # 记下每轮对话speaker1开始的时间戳
        if item['speaker'] == '1':
            begin = int(item['bg'])

            detect_gap(begin, end)

            speaker1_quiet_time += (begin - end) / 1000

            item_s1['speaker'] = '1'
            item_s1['begin'] = time_format(int(item['bg']))  # 转换成 分:秒格式,例如：01:12

            str_s1 = ''
            # 循环读到speaker2为止
            while item['speaker'] == '1':
                str_s1 += item['onebest']  # 把同一个说话者一次说的话放在一起，直到另一说话者说话时为止
                item_s1['end'] = time_format(int(item['ed']))

                # 把当前讲话者讲的内容和时间戳加入列表
                raw_dialog_dict['begin'] = time_format(int(item['bg']))
                raw_dialog_dict['end'] = time_format(int(item['ed']))
                raw_dialog_dict['speaker'] = 'speaker1'
                raw_dialog_dict['text'] = item['onebest']

                # print(raw_dialog_dict)

                raw_dialog_list.append(raw_dialog_dict)
                # print(raw_dialog_list)

                end = int(item['ed'])
                t = (end - int(item['bg'])) / 1000  # 只计算说话秒数
                # print("speaker1's time:'", t)
                speaker1_time += t

                # 同一位讲话者，间隔一下之后继续说的话
                index += 1
                if index < length:
                    item = raw_list[index]  # 换下一段话
                    if item['speaker'] == '1':  # 同一讲话者
                        begin = int(item['bg'])
                        # print('speaker1:', time_format(int(item['bg'])))
                        # print('%d-%d=%d'%(begin, end, (begin-end)//1000))
                        # detect_gap(begin, end)    #用这个不能正常显示
                        gap = (begin - end) / 1000
                        # print('speaker1 gap:', gap)
                        speaker1_quiet_time += gap

                        end = int(item['ed'])  # new end                
                        gap = round(gap)
                        if gap >= PAUSE_MIN_SECONDS:  # 安静指定秒以上才显示
                            # str_s1 +=  '【安静%d秒】' %(gap) #不用安静起始的时间戳，文本看起来会清爽一些
                            str_s1 += item_s1['end'] + '【安静%d秒】' % (gap)  # 应该直接加到当前讲话者的内容里
                    else:
                        break
                else:
                    break
            item_s1['text'] = str_s1
            # print("item_s1:", item_s1)

            text = "%s  %s：%s\n" % (item_s1['begin'], speaker1, str_s1)
            # total_length_speaker1 += len(text)
            total_length_speaker1 += count_chars(text)

            speaker1_content = ''.join([speaker1_content, text])
            file_content = ''.join([file_content, text])

            # end = int(item['ed'])
        # 记下每轮对话speaker2开始的时间戳
        if item['speaker'] == '2':
            begin = int(item['bg'])

            detect_gap(begin, end)
            speaker2_quiet_time += (begin - end) / 1000

            item_s2['speaker'] = '2'
            item_s2['begin'] = time_format(int(item['bg']))
            str_s2 = ''
            # 循环读到speaker1为止
            while item['speaker'] == '2':
                str_s2 += item['onebest']
                item_s2['end'] = time_format(int(item['ed']))

                # 把当前讲话者讲的内容和时间戳加入列表
                raw_dialog_dict['begin'] = time_format(int(item['bg']))
                raw_dialog_dict['speaker'] = 'speaker2'
                raw_dialog_dict['end'] = time_format(int(item['ed']))
                raw_dialog_dict['text'] = item['onebest']

                if index < 5:
                    print(raw_dialog_dict)
                    raw_dialog_list.append(raw_dialog_dict)
                    print(raw_dialog_list)

                end = int(item['ed'])
                t = (end - int(item['bg'])) / 1000  # 只计算说话秒数
                # print("speaker2's time:'", t)
                speaker2_time += t

                index += 1
                if index < length:
                    item = raw_list[index]
                    if item['speaker'] == '2':
                        begin = int(item['bg'])
                        # detect_gap(begin, end)
                        gap = (begin - end) / 1000
                        # print('speaker2 gap:', gap)
                        speaker2_quiet_time += gap
                        end = int(item['ed'])  # new end
                        gap = round(gap)
                        if gap >= PAUSE_MIN_SECONDS:  # 安静指定秒以上才显示
                            str_s2 += item_s2['end'] + '【安静%d秒】' % (gap)  # 应该直接加到当前讲话者的内容里
                    else:
                        break
                else:
                    break
            item_s2['text'] = str_s2
            text = "%s  %s：%s\n" % (item_s2['begin'], speaker2, str_s2)
            # total_length_speaker2 += len(text)
            total_length_speaker2 += count_chars(text)

            speaker2_content = ''.join([speaker2_content, text])
            file_content = ''.join([file_content, text])

    total_time = time_format(int(item['ed']))

    # 调用函数，编辑讯飞语音中的常见错误：
    speaker1_content = edit_common_error(speaker1_content)
    speaker2_content = edit_common_error(speaker2_content)
    file_content = edit_common_error(file_content)

    # 统计重要数据：说话时长、安静时长、字节数及相应占比
    def calc_statistics():

        # 计算两说话时长占比
        speaker1_time_ratio = speaker1_time / (speaker1_time + speaker2_time)
        # fw.write('speaker1:' + ' ' * 10 + '说话时长：%d秒,占比：%.2f%%  安静时长:%d 秒。\n' % (speaker1_time, speaker1_time_ratio * 100, output_result[3]))
        # fw.write('speaker2:' + ' ' * 10 + '说话时长：%d秒,占比：%.2f%%  安静时长:%d 秒。\n' % (speaker2_time, (1-speaker1_time_ratio) * 100, output_result[4]))

        # 计算字节两人说话占比
        speaker1_words_ratio = total_length_speaker1 / (total_length_speaker1 + total_length_speaker2)
        # fw.write('speaker1:' + ' ' * 10 + '字节总数：%d,占比：%.2f%%\n' % (total_length_speaker1, ratio * 100))
        # fw.write('speaker2:' + ' ' * 10 + '字节总数：%d,占比：%.2f%%\n' % (total_length_speaker2, (1 - ratio) * 100))

        speaker1_dict = OrderedDict()
        speaker1_dict["说话者"] = "speker1"
        speaker1_dict["说话时长"] = str(round(speaker1_time)) + "秒"
        speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
        speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
        speaker1_dict["字符总数"] = "%d" % total_length_speaker1
        speaker1_dict["字符占比"] = "%.2f%%\n" % (speaker1_words_ratio * 100)

        speaker2_dict = OrderedDict()
        speaker2_dict["说话者"] = "speker2"
        speaker2_dict["说话时长"] = str(round(speaker2_time)) + "秒"
        speaker2_dict["时长占比"] = "%.2f%%" % ((1 - speaker1_time_ratio) * 100)
        speaker2_dict["安静时长"] = "%d 秒" % speaker2_quiet_time
        speaker2_dict["字符总数"] = "%d" % total_length_speaker2
        speaker2_dict["字符占比"] = "%.2f%%\n" % ((1 - speaker1_words_ratio) * 100)

        # print("speaker1:", speaker1_dict)
        # print("speaker2:", speaker2_dict)

        return speaker1_dict, speaker2_dict

    # print(raw_dialog_list)

    speaker1_dict, speaker2_dict = calc_statistics()

    result = (file_content, speaker1_content, speaker2_content, speaker1_dict, speaker2_dict, raw_dialog_list)

    return result


# raw,只按原始转录内容，加上时间戳
# 用途：方便用户对照着进行编辑，以纠正原始转录时可能出现的错误，例如：特别是搞错了是谁说的话
# 需要先调用convert_data函数,
# output_result是由函数convert_data返回的元组
# 默认是输出两位说话者的文字two， 可选择speaker1, or speaker2，只输出教练的文字
def write_raw_to_txt_file(output_file):
    raw_list = get_result_list(data_file)  # 获取讯飞原始转写的数据信息

    try:
        fw = open(output_file, 'w', encoding='utf-8')
        all_text = ''
        for item in raw_list:
            begin = time_format(int(item['bg']))  # 按时间戳显示开始时间
            # 处理一下原始转写的文字
            text = edit_common_error(item['onebest'])
            # 不处理，保留原始转写的结果
            text = item['onebest']
            text = begin + '  speaker' + item['speaker'] + '：' + text + '\n'  # 不加换行，否则太多！
            # print(text)
            all_text = ''.join([all_text, text])
        fw.write(all_text)
        print('从原始文件转成时间戳文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 只输出文本，不显示时间戳，只按1个讲话者
def write_raw_without_stamp_to_txt_file(output_file):
    raw_list = get_result_list(data_file)  # 获取讯飞原始转写的数据信息

    try:
        fw = open(output_file, 'w', encoding='utf-8')
        text = ''
        for item in raw_list:
            # begin = time_format(int(item['bg']))  # 按时间戳显示开始时间
            # 处理一下原始转写的文字
            text = ''.join([text, edit_common_error(item['onebest'])])
            # 不处理，保留原始转写的结果
            # print(text)

        fw.write(text)
        print('从原始文件转成时间戳文件:%s成功！' % output_file)
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 文件开头、统计字节数及笔录写入文件
# 需要先调用convert_data函数
# output_result是由函数convert_data返回的元组
# 默认是输出两位说话者的文字two， 可选择speaker1, or speaker2，只输出教练的文字
def write_to_txt_file(output_file, output_result, speaker="two"):
    try:
        fw = open(output_file, 'w', encoding='utf-8')
        # 写文件开头
        fw.write('教练对话语音转写\n\n')
        fw.write('注：请根据实际情况，替换speaker1或speaker2为教练或客户：\n')
        fw.write('总时长：%s\n' % total_time)  # total_time is calculated in function: convert_data

        # 打印标题统计speaker1, speaker2的讲话及安静信息：
        text = ''
        speaker1_dict = output_result[3]
        length = 0
        for key in speaker1_dict:
            value = speaker1_dict[key]
            text = ''.join([text, key, value])
            length += 1
            if length < len(speaker1_dict):
                text += '，'

        fw.write(text + '\n')

        text = ''
        speaker2_dict = output_result[4]
        length = 0
        for key in speaker2_dict:
            value = speaker2_dict[key]
            text = ''.join([text, key, value])
            length += 1
            if length < len(speaker2_dict):
                text += '，'
        fw.write(text + '\n')

        fw.write('笔录:\n')

        # 写笔录
        if speaker == "two":  # 教练和客户的对话
            fw.write(output_result[0])
        elif speaker == "speaker1":  # 仅一方的讲话，根据参数决定
            fw.write(output_result[1])
        elif speaker == "speaker2":
            fw.write(output_result[2])
        print("保存结果到{}成功。".format(output_file))
    except IOError:
        print("保存失败!")
    finally:
        fw.close()


# 需要先调用convert_data函数
# 自动生成word文档，包括:标题、统计信息、对话时间戳及对话内容（教练的话以蓝色字体标记）
def write_to_word_file(output_file, output_result, speaker="two", coach="speaker1"):
    # 文件名称
    file_name = output_file
    # 打开文件
    document = docx.Document()

    # 添加指定格式的段落
    # text：段落内容
    def add_formated_text(document, text, font_cn_name="宋体", size=12, color=RGBColor(0, 0, 0), alignment="left",
                          is_bold=False):
        # 添加段落
        p1 = document.add_paragraph()
        text1 = p1.add_run(text)

        text1.font.name = 'Times New Roman'  # 控制是西文时的字体
        text1.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn_name)  # 控制是中文时的字体

        text1.font.size = Pt(size)
        text1.font.color.rgb = color

        # 此处高亮显示“安静”为绿色， 如果有的话
        # 此处还需改进。。。教练的文字，字体不对，颜色蓝色也没有了
        # highlight_keywords(p1, text, keyword, color)

        # 加粗
        text1.font.bold = is_bold
        # 斜体
        # text1.font.italic = True

        if alignment == "center":
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "left":
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p1.paragraph_format.space_after = Pt(10)  # 段后间距，default

        document.save(file_name)

    # 第一段（设置标题居中）
    title = '教练对话语音转写'
    add_formated_text(document, title, font_cn_name="宋体", size=16, color=RGBColor(0, 0, 255), alignment="center",
                      is_bold=True)

    # remind ='注：请根据实际情况，替换speaker1或speaker2为教练或客户。'
    # add_formated_text(document, remind)

    total_time = '总时长：59分钟'
    add_formated_text(document, total_time)

    # 添加表格, 显示统计教练和客户的信息
    table = document.add_table(rows=3, cols=6,
                               style="Light List Accent 5")  # 指定表格格式，参见：https://blog.csdn.net/ibiao/article/details/78595295
    table.style.font.name = '宋体'  # 表格内的字体类型
    table.style.font.size = 120000  # 字号

    # 在表的指定行，列，输入内容
    def add_text2table(table, row, col, text):
        cell = table.cell(row, col)
        cell.text = text

    # 需要用到排序字典，#循环遍历，添加入表格：
    speaker1_dict = output_result[3]
    speaker2_dict = output_result[4]

    row, col = 0, 0
    # 表格各列标题
    for key in speaker1_dict:
        add_text2table(table, row, col, key)
        col += 1

    row, col = 1, 0
    for key in speaker1_dict:
        if col == 0:  # 根据参数决定哪位是教练，默认第一位是教练
            if coach == "speaker1":
                value = "教练"
            else:
                value = "客户"
        else:
            value = speaker1_dict[key]
            # print(key, value)
        add_text2table(table, row, col, value)
        col += 1

    row, col = 2, 0
    for key in speaker2_dict:
        if col == 0:  # 根据参数决定哪位是教练，默认第一位是教练
            if coach == "speaker2":
                value = "教练"
            else:
                value = "客户"
        else:
            value = speaker2_dict[key]
            # print(key, value)
        add_text2table(table, row, col, value)
        col += 1

    text = "笔录:"
    add_formated_text(document, text)

    # 获取原始数据，转换成列表，以便操作
    speaking_list = output_result[0].split('\n')
    for item in speaking_list:
        if "speaker1" in item:
            if speaker == "speaker2":  # 跳过，不输出
                continue
            if coach == "speaker1":
                # 教练文字用蓝色、加粗
                item = item.replace("speaker1", "教练")

                add_formated_text(document, item, color=RGBColor(0, 0, 255), is_bold=True)

                # print("coach:", item)
            else:
                item = item.replace("speaker1", "客户")
                add_formated_text(document, item)
        elif "speaker2" in item:
            if speaker == "speaker1":  # 跳过，不输出
                continue
            if coach == "speaker2":
                # 教练文字用蓝色、加粗
                item = item.replace("speaker2", "教练")
                add_formated_text(document, item, color=RGBColor(0, 0, 255), is_bold=True)
                # print("coach:", item)
            else:
                item = item.replace("speaker2", "客户")
                add_formated_text(document, item)

    print("保存结果到{}成功。".format(output_file))


# 按照讯飞的转写结果显示，可以让用户手动编辑（删除多余的时间戳），
# 避免因讯飞识别发音人出错而导致时间戳错误
def get_result_list(filename):
    try:
        fr = open(filename, 'r', encoding='utf-8')

        str_data = fr.read()
        data_dict = eval(str_data)  # 字符串转换为字典类型
        # print(data_dict)
        result_list = json.loads(data_dict['data'])  # 将json格式的字符串转为列表
        # print(result_list)
        print("从文件{}读取数据成功。".format(filename))
        return result_list
    except:
        print("从文件{}读取数据失败！".format(filename))


# 把按轮流对话的方式存入文件，字典的列表（字符串）格式
def save_to_file(filename):
    try:
        fw = open(filename, 'w', encoding='utf-8')
        fw.write(str(convert_data(data_file)))
        print("保存结果到{}成功。".format(filename))
    except:
        print("保存失败!")
    finally:
        fw.close()


if __name__ == '__main__':
    # data_file = 'JetCoachDeng-Sample-1m.txt'
    # output_file = 'JetCoachDeng-Sample-1m-按speakers整理.txt'

    # data_file = 'jet coach mengyao-sample.txt' #讯飞语音已经转写好的文件
    # data_file = '../output/' + "2020-0605-教练梦瑶-画大饼的思维.mp3-讯飞转写原始结果.txt"

    audio_file = r"C:\Jet's Docs\Jet 常用文档-Asus\#1-STEAM编程教育\#1-大陆JH学校编程培训项目\#1-编程培训项目文档\2020-Moodle在线课程录制\Introduction\0.0-课程介绍v.1.mp3"
    data_file = audio_file + "-讯飞转写原始结果.txt"

    output_file = audio_file + "-纯文字结果.txt"
    write_raw_without_stamp_to_txt_file(output_file)

    '''
    output_result = convert_data(data_file)
    #输出成文本文件
    output_file = data_file + '-raw_教练和客户.txt'
    write_raw_to_txt_file(output_file)
    '''

    '''
    output_file = data_file + '-整理_教练和客户.txt'
    write_to_txt_file(output_file, output_result, speaker='two')
    
    output_file = data_file + '-整理_教练only.txt'
    write_to_txt_file(output_file, output_result, speaker='speaker1')
    
    output_file = data_file + '-整理_客户only.txt'
    write_to_txt_file(output_file, output_result, speaker='speaker2')
    
    #输出成word文件
    output_word = data_file + '-整理_教练和客户.doc'
    write_to_word_file(output_word, output_result, speaker='two', coach='speaker1')
    
    output_word = data_file + '-整理_教练only.doc'
    write_to_word_file(output_word, output_result, speaker='speaker1', coach='speaker1')
    
    '''

    '''
    output_word = data_file + '-整理_客户only.doc'
    write_to_word_file(output_word, output_result, speaker='speaker2', coach='speaker2')
    '''

    # save_to_file(output_file)
