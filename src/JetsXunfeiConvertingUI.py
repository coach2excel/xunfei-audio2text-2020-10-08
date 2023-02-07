# -*- coding: utf-8 -*-

"""
@author Jet
@desc 讯飞转写的主界面
@date 2020/06/20

#注：用VSCode运行，上传语音文件，等待转写到一段时间就会退出程序！

"""
import pickle
import sys
import os
import time

import pygame
from PyQt5.QtCore import QUrl
from PyQt5.QtGui import QFont, QBrush, QColor
from PyQt5.QtMultimedia import QMediaPlayer, QMediaContent  # 支持mp3, mp4, wav格式播放

from PyQt5.QtWidgets import QMainWindow, QApplication, QFileDialog, QMessageBox, QDialog, QTableWidgetItem, QStyle, \
    QSlider, QHeaderView, QAbstractItemView
from PyQt5 import QtCore, QtGui, QtWidgets

# 以下生成word文档需要
from docx.oxml import OxmlElement, ns

'''
Jet: 
pip install python-docx

不能安装pip install docx,否则运行时会提示错误：
  File "C:\lib\site-packages\docx.py", line 30, in <module>
    from exceptions import PendingDeprecationWarning
ModuleNotFoundError: No module named 'exceptions'
 if already installed: pip unistall docx


'''
import docx
from docx.shared import Inches, Cm, Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import qn  # 字体设置
# 修改word关键词颜色
from docx.oxml.ns import qn  # 设置中文字体

# import subprocess #用于显示调用讯飞过程中在cmd中打印的信息
import vlc  # 播放声音 pip install python-vlc
import VLC_Jet  # Jet 改编的VLC API  # 会导入上一行的pip install python-vlc
import getaudioinfo
from convert_audio_progress_gui import ConvertingDialog
import transform_tool

'''
修改记录：
06-28 从qt_coaching4_modify_to_Run_4.6.py另存为本文件
2020-06-13 添加按钮：提交讯飞原始转换
2020-06-12 实现功能：radiobutton选择教练/客户

'''

# Jets: 主窗口
"""
class ParentWindow(QMainWindow):
    def __init__(self):
        QMainWindow.__init__(self)
        self.main_ui = XunFeiConvertingUIMainWindow()
"""

# Jets: 子窗口，讯飞转写监控（显示进度和提示信息）
"""
class ChildWindow(QDialog):
    # Jet： 用户选择是1位还是2位讲话者
    def __init__(self, audio_file, raw_txt_file, speaker_num=2):
        QDialog.__init__(self)
        self.child = ConvertingDialog(audio_file, raw_txt_file, speaker_num)  # Jet: 显示转写进度的窗口
        #self.child.setupUi(self)
        # 在主窗口已经选择好了
        #self.audio_file = audio_file
        #self.raw_txt_file = raw_txt_file
        #self.child.xunfei_converting(self.audio_file, self.raw_txt_file)
"""
"""
class ChildWindow(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("打开子窗口")
"""


# Jets：从ui文件生成py文件之后，需要作相应的修改（见相应注释）
# class Ui_MainWindow(object):#ui自动生成的
class XunFeiConvertingUIMainWindow(QMainWindow):  # 手工修改成这样
    """
    Jets讯飞转写的主界面
    """

    # Jet：这是需要增加的
    def __init__(self):
        super().__init__()

        # 指定播放时间很大差距，放弃使用！
        self.mediaPlayer = QMediaPlayer(None, QMediaPlayer.VideoSurface)  # 用于播放音频

        self.audio_file = None  # 要转写和播放的文件wav/mp3
        self.audio_time = 0  # 音频时长，秒数，在选择好音频文件后再计算出来

        self.vlc_player = VLC_Jet.VlcPlayer()  # 用于播放音频

        # pygame.init()  #为播放音乐做准备
        # pygame.mixer.init() #播放mp3需要
        # self.track = None
        # 不用pygame了，改用QMediaPlayer, 支持mp3, mp4, wav格式

        self.initUI()  # Jet：这是需要增加的,在下面编写具体代码

        self.output_timestampbinary_file = None  # 要转写成二进制文件的全路径名
        self.stamp_dict_list = None  # 存储从二进制文件中读出的字典列表

        self.speaker1 = "教练"  # 默认speaker1
        self.speaker2 = "客户"

        self.stamp_list = None #GUI左下方显示的数据来源（从raw.txt转换的，带起止时间戳，教练/客户）
        self.icf_list = None  # ICF格式的教练/客户对话（带起止时间戳）
        self.icf_stat_list = None  # ICF教练对话的统计信息
        self.is_stat_included = True  # 默认输出包含统计信息
        self.is_only_coach_text = False  # 默认不是只输出教练说的文字
        self.icf_list_cache = None  # 用于缓存ICF list table widget 中的数据
        self.icf_stat_cache = None  # 用于缓存ICF stat table widget 中的数据
        self.coaching_row_count = 0  # 合并讲话后，教练的讲话条数（行数）
        self.raw_output_result = None  # 缓存从原始转写raw.txt获取的数据
        # 来自 get_stamp_list_from_rawtxt调用了convert_2speakers_rawtxt
        # in transform_tool.py

        # Jet: 用户选择是1个还是2个讲话者

    def onRadioButton_SpeakerNum_Clicked(self):
        radioBtn = self.sender()
        if radioBtn.isChecked():
            if radioBtn.text() == "1位讲话者":
                # print("speaker number is 1")
                self.speaker_num = 1
                self.radioButton_2speakers.setChecked(False)

            if radioBtn.text() == "2位讲话者":
                # print("speaker number is 2")
                self.speaker_num = 2
                self.radioButton_1speaker.setChecked(False)

            self.statusbar.showMessage("选择的语音中有%d位讲话者:" % self.speaker_num)

    # 根据用户选择Speaker1的Radiobutton, 确定speaker1和speaker2的身分
    def onRadioButtonClicked(self):
        radioBtn = self.sender()
        if radioBtn.isChecked():
            if radioBtn.text() == "教练":
                print("用户选择Speaker1的Radiobutton: speaker1 is:教练")
                self.speaker1 = "教练"
                self.speaker2 = "客户"
                self.statusbar.showMessage("speaker1 is:" + self.speaker1)
                self.radioButton_s2iscoach.setChecked(False)
                self.radioButton_s2isclient.setChecked(True)

                # 2020-07-20 刷新显示
                # self.shiftspeakers_in_rawstamplist_table()  # 运行很慢，987条数据要花30-40秒！
                # to do : fix above bug
                self.update_display()


            if radioBtn.text() == "客户":
                print("用户选择Speaker1的Radiobutton: speaker1 is:客户")

                self.speaker1 = "客户"
                self.speaker2 = "教练"
                self.statusbar.showMessage("speaker1 is:" + self.speaker1)
                self.radioButton_s2iscoach.setChecked(True)
                self.radioButton_s2isclient.setChecked(False)

                # 2020-07-20 刷新显示
                # self.shiftspeakers_in_rawstamplist_table() # 运行很慢，987条数据要花30-40秒！
                # to do : 解决方案：分页显示
                # 一、想用QTabelWidge的话就实现类似网页中上一页 / 下一页的效果，一次只加载一部分数据。毕竟，用户不用一次看很多条
                #  二、不用QTableWidget，用QTableView和model代替
                # to do : fix above bug

                self.update_display()

    # adds on  2020/7
    def update_display(self):
        self.switch_coach_client_in_stamp_list()  # 互换内存中教练/客户,左侧表格显示的数据来源
        self.switch_coach_client_in_stat()  # 互换内存中教练/客户的统计信息
        self.switch_coach_client_in_icf_list()  # 互换内存中讲话者：教练/客户

        #self.display_stamplist_in_table()
        # self.display_2speakers_rawstamplist_in_table()  上千条数据更新缓慢，不刷新了，以后修改此函数为分页显示即可。
        self.display_icf_format()
        self.display_icfstat_in_table()

    # 根据用户选择Speaker2的Radiobutton, 确定speaker1和speaker2的身分
    def onRadioButtonClicked2(self):
        radioBtn = self.sender()
        if radioBtn.isChecked():
            if radioBtn.text() == "教练":
                print("speaker1 is:客户")
                self.speaker1 = "客户"
                self.statusbar.showMessage("speaker1 is:" + self.speaker1)
                self.radioButton_s1iscoach.setChecked(False)
                self.radioButton_s1isclient.setChecked(True)

                # 2020-07-20 刷新显示
                # self.display_2speakers_rawstamplist_in_table()  # 有Bug,会死机！
                # self.display_icf_format(self.raw_txt_file)
                # self.display_icfstat_in_table()

            if radioBtn.text() == "客户":
                print("speaker1 is:教练")
                self.speaker1 = "教练"
                self.statusbar.showMessage("speaker1 is:" + self.speaker1)
                self.radioButton_s1iscoach.setChecked(True)
                self.radioButton_s1isclient.setChecked(False)

                # 2020-07-20 刷新显示
                # self.display_2speakers_rawstamplist_in_table()  # 有Bug,会死机！
                # self.display_icf_format(self.raw_txt_file)
                # self.display_icfstat_in_table()

    # Jets: 生成界面。需要修改：ui默认生成的是 def setupUi(self, self):
    def initUI(self):
        # Jet：需要把MainWindow替代为self
        # MainWindow.setObjectName("MainWindow")
        # MainWindow.resize(1284, 982)
        # self.centralwidget = QtWidgets.QWidget(MainWindow)

        self.setObjectName("MainWindow")
        self.resize(1400, 1250)

        # 初始化要选择的音频文件和要保存的原始转写文件
        self.audio_file = ""
        self.raw_txt_file = ""
        self.speaker_num = 2  # 默认音频文件是2个讲话者
        # self.cwd = os.getcwd()  # 获取当前程序文件位置

        self.centralwidget = QtWidgets.QWidget(self)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayoutWidget = QtWidgets.QWidget(self.centralwidget)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(10, 10, 1380, 1220))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")

        self.gridLayout = QtWidgets.QGridLayout(self.gridLayoutWidget)
        self.gridLayout.setContentsMargins(10, 10, 10, 10)
        self.gridLayout.setSpacing(10)
        self.gridLayout.setObjectName("gridLayout")

        # 第一行的控件
        self.btn_locate_audio_file = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_locate_audio_file.setObjectName("btn_locate_audio_file")
        self.gridLayout.addWidget(self.btn_locate_audio_file, 0, 0, 1, 1)
        """
        int fromRow：所在行
        int fromColumn：所在列
        int rowSpan：所占行数
        int columnSpan：所占列数
        """

        self.lineEdit_audio_file = QtWidgets.QLineEdit(self.gridLayoutWidget)
        self.lineEdit_audio_file.setObjectName("lineEdit_audio_file")
        self.gridLayout.addWidget(self.lineEdit_audio_file, 0, 1, 1, 3)

        # 对RadioButton分组: 选择是1个还是2个讲话者
        self.radioButtonGroup0 = QtWidgets.QButtonGroup(self.gridLayoutWidget)

        self.radioButton_1speaker = QtWidgets.QRadioButton(self.gridLayoutWidget)
        self.radioButtonGroup0.addButton(self.radioButton_1speaker, 1)
        self.radioButton_1speaker.setChecked(False)
        self.radioButton_1speaker.setObjectName("radioButton_1speaker")
        self.gridLayout.addWidget(self.radioButton_1speaker, 0, 4, 1, 1)

        self.radioButton_2speakers = QtWidgets.QRadioButton(self.gridLayoutWidget)
        self.radioButtonGroup0.addButton(self.radioButton_2speakers, 1)
        self.radioButton_2speakers.setChecked(True)
        self.radioButton_2speakers.setObjectName("radioButton_2speakers")
        self.gridLayout.addWidget(self.radioButton_2speakers, 0, 5, 1, 1)

        # 第二行的控件
        # 此按钮：打开对话框，设置原始转换的结果另存为的文件名（全路径+默认文件名）
        self.btn_set_raw_output_file = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_set_raw_output_file.setObjectName("btn_set_raw_output_file")
        # Jet:参数说明：gridLayout.addWidget(widget, row, column, columnspan, alignment)
        self.gridLayout.addWidget(self.btn_set_raw_output_file, 1, 0, 1, 1)

        self.lineEdit_rawtxt_file = QtWidgets.QLineEdit(self.gridLayoutWidget)
        self.lineEdit_rawtxt_file.setObjectName("lineEdit_rawtxt_file")
        self.gridLayout.addWidget(self.lineEdit_rawtxt_file, 1, 1, 1, 3)

        # btn提交讯飞原始转换
        self.btn_xunfei2txt = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_xunfei2txt.setObjectName("btn_xunfei2txt")
        self.gridLayout.addWidget(self.btn_xunfei2txt, 1, 4, 1, 2)

        # 第三行的控件
        self.label = QtWidgets.QLabel(self.gridLayoutWidget)
        self.label.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.label.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 2, 0, 1, 1)

        # 对RadioButton分组
        self.radioButtonGroup1 = QtWidgets.QButtonGroup(self.gridLayoutWidget)
        self.radioButton_s1iscoach = QtWidgets.QRadioButton(self.gridLayoutWidget)
        self.radioButtonGroup1.addButton(self.radioButton_s1iscoach, 1)

        self.radioButton_s1iscoach.setChecked(True)
        self.radioButton_s1iscoach.setObjectName("radioButton_s1iscoach")

        self.gridLayout.addWidget(self.radioButton_s1iscoach, 2, 1, 1, 1)

        # self.radioButton_s1iscoach.toggled.connect(self.onRadioButtonClicked)#Jet:注意，函数调用不能加()

        self.radioButton_s1isclient = QtWidgets.QRadioButton(self.gridLayoutWidget)
        # self.radioButton_s1isclient.setChecked(False)
        # self.radioButton_s1isclient.setCheckable(True)
        self.radioButton_s1isclient.setObjectName("radioButton_s1isclient")
        self.gridLayout.addWidget(self.radioButton_s1isclient, 2, 2, 1, 1)
        self.radioButtonGroup1.addButton(self.radioButton_s1isclient, 2)

        # self.radioButton_s1isclient.toggled.connect(self.onRadioButtonClicked)

        self.label_2 = QtWidgets.QLabel(self.gridLayoutWidget)
        self.label_2.setAlignment(QtCore.Qt.AlignRight | QtCore.Qt.AlignTrailing | QtCore.Qt.AlignVCenter)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 2, 3, 1, 1)

        # 对RadioButton分组
        self.radioButtonGroup2 = QtWidgets.QButtonGroup(self.gridLayoutWidget)

        self.radioButton_s2isclient = QtWidgets.QRadioButton(self.gridLayoutWidget)
        # self.radioButton_s2isclient.setCheckable(False)
        self.radioButton_s2isclient.setChecked(True)
        self.radioButton_s2isclient.setObjectName("radioButton_s2isclient")
        self.gridLayout.addWidget(self.radioButton_s2isclient, 2, 4, 1, 1)

        # self.radioButton_s2isclient.toggled.connect(self.onRadioButtonClicked2)

        self.radioButton_s2iscoach = QtWidgets.QRadioButton(self.gridLayoutWidget)
        # self.radioButton_s2iscoach.setCheckable(False)
        self.radioButton_s2iscoach.setChecked(False)
        self.radioButton_s2iscoach.setObjectName("radioButton_s2iscoach")
        self.gridLayout.addWidget(self.radioButton_s2iscoach, 2, 5, 1, 1)
        self.radioButtonGroup2.addButton(self.radioButton_s2iscoach, 1)
        self.radioButtonGroup2.addButton(self.radioButton_s2isclient, 2)

        # self.radioButton_s2iscoach.toggled.connect(self.onRadioButtonClicked2)
        # 第4行
        self.label_3 = QtWidgets.QLabel(self.gridLayoutWidget)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 3, 0, 1, 1)
        self.label_4 = QtWidgets.QLabel(self.gridLayoutWidget)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 3, 3, 1, 1)

        '''
        #不再显示
        self.label_5 = QtWidgets.QLabel(self.gridLayoutWidget)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 3, 2, 1, 2)
        '''
        self.label_6 = QtWidgets.QLabel(self.gridLayoutWidget)
        self.label_6.setObjectName("label_6")
        self.gridLayout.addWidget(self.label_6, 3, 4, 1, 2)

        # 第5行
        self.btn_import_rawtxt = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_import_rawtxt.setObjectName("btn_import_rawtxt")
        self.gridLayout.addWidget(self.btn_import_rawtxt, 4, 0, 1, 1)

        ''' 以下按钮弃用  2020/07/09
        self.btn_save_editing2stampfile = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_save_editing2stampfile.setObjectName("btn_save_raw_to_txt")
        self.gridLayout.addWidget(self.btn_save_editing2stampfile, 4, 1, 1, 1)

        self.btn_import_stampbinary = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_import_stampbinary.setObjectName("btn_import_stampbinary")
        self.gridLayout.addWidget(self.btn_import_stampbinary, 4, 2, 1, 1)
        
        self.btn_display_icf_format = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_display_icf_format.setObjectName("btn_display_icf_format")
        self.gridLayout.addWidget(self.btn_display_icf_format, 4, 3, 1, 1)
        '''

        # 添加按钮 2020/07/09
        # 批量编辑原始文字
        self.btn_batch_edit = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_batch_edit.setObjectName("btn_batch_edit")
        self.gridLayout.addWidget(self.btn_batch_edit, 4, 2, 1, 1)

        self.btn_save_icf_to_binary = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_save_icf_to_binary.setObjectName("btn_save_icf_to_binary")
        self.gridLayout.addWidget(self.btn_save_icf_to_binary, 4, 4, 1, 1)

        self.btn_import_icf_from_binary = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_import_icf_from_binary.setObjectName("btn_import_icf_from_binary")
        self.gridLayout.addWidget(self.btn_import_icf_from_binary, 4, 5, 1, 1)

        self.ckbox_export_stat = QtWidgets.QCheckBox(self.gridLayoutWidget)
        self.ckbox_export_stat.setObjectName("ckbox_export_stat")
        self.gridLayout.addWidget(self.ckbox_export_stat, 4, 3, 1, 1)
        self.ckbox_export_stat.setChecked(True)

        # 第6行
        # added 2020/07/09 只保留教练文字
        self.ckbox_only_coach_text = QtWidgets.QCheckBox(self.gridLayoutWidget)
        self.ckbox_only_coach_text.setObjectName("ckbox_export_stat")
        self.gridLayout.addWidget(self.ckbox_only_coach_text, 5, 3, 1, 1)
        self.ckbox_only_coach_text.setChecked(False)

        self.btn_save_icf_to_txt = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_save_icf_to_txt.setObjectName("btn_save_icf_to_txt")
        self.gridLayout.addWidget(self.btn_save_icf_to_txt, 5, 4, 1, 1)

        self.btn_save_icf_to_word = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_save_icf_to_word.setObjectName("btn_save_icf_to_word")
        self.gridLayout.addWidget(self.btn_save_icf_to_word, 5, 5, 1, 1)

        # 第7行
        # 显示：Jet根据讯飞原始txt转写成的起止时间戳格式（包括显示安静的内容）
        self.tableWidget_raw = QtWidgets.QTableWidget(self.gridLayoutWidget)
        self.tableWidget_raw.setObjectName("tableWidget_raw")
        self.tableWidget_raw.setColumnCount(0)
        self.tableWidget_raw.setRowCount(0)
        self.gridLayout.addWidget(self.tableWidget_raw, 7, 0, 1, 3)

        # 显示ICF对话格式：同一位讲话者的一段话汇兑在一起显示（包括显示安静的内容）
        self.tableWidget_icf = QtWidgets.QTableWidget(self.gridLayoutWidget)
        self.tableWidget_icf.setObjectName("tableWidget_icf")
        self.tableWidget_icf.setColumnCount(0)
        self.tableWidget_icf.setRowCount(0)
        self.gridLayout.addWidget(self.tableWidget_icf, 7, 3, 1, 3)

        # ICF统计信息
        self.tableWidget_stat = QtWidgets.QTableWidget(self.gridLayoutWidget)
        self.tableWidget_stat.setObjectName("tableWidget_stat")
        self.tableWidget_stat.setColumnCount(0)
        self.tableWidget_stat.setRowCount(0)
        self.gridLayout.addWidget(self.tableWidget_stat, 8, 3, 3, 3)

        # 音频播放控件----------------
        self.slider_audio = QtWidgets.QSlider(self.gridLayoutWidget)
        self.slider_audio.setOrientation(QtCore.Qt.Horizontal)
        self.slider_audio.setObjectName("slider_audio")
        self.slider_audio.setRange(0, 0)

        self.gridLayout.addWidget(self.slider_audio, 8, 0, 1, 3)

        # 用不上了
        '''
        self.btn_stop_audio = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_stop_audio.setObjectName("btn_stop_audio")
        self.gridLayout.addWidget(self.btn_stop_audio, 9, 1, 1, 1)
        '''
        self.btn_play_audio = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.btn_play_audio.setObjectName("btn_play_audio")
        self.btn_play_audio.setEnabled(False)
        self.btn_play_audio.setIcon(self.style().standardIcon(QStyle.SP_MediaPlay))
        self.gridLayout.addWidget(self.btn_play_audio, 9, 0, 1, 1)

        # ---------------------

        # 设置列在水平方向的空间比例分配
        self.gridLayout.setColumnStretch(0, 1)
        self.gridLayout.setColumnStretch(1, 1)
        self.gridLayout.setColumnStretch(2, 1)
        self.gridLayout.setColumnStretch(3, 1)
        self.gridLayout.setColumnStretch(4, 1)
        self.gridLayout.setColumnStretch(5, 1)

        # 设置行在垂直方向的空间比例分配
        # self.gridLayout.setRowStretch(7, 10)  # ICF tableWidget高度占比
        # self.gridLayout.setRowStretch(8, 8)  # ICF stat tableWidget高度占比
        self.setCentralWidget(self.centralwidget)

        ''' #取消菜单栏
        self.menubar = QtWidgets.QMenuBar(self)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 1284, 23))
        self.menubar.setObjectName("menubar")
        self.menu = QtWidgets.QMenu(self.menubar)
        self.menu.setObjectName("menu")
        self.setMenuBar(self.menubar)
        self.menubar.addAction(self.menu.menuAction())
        '''
        # Jet: 添加状态栏
        self.statusbar = QtWidgets.QStatusBar(self)
        self.statusbar.setEnabled(True)
        # self.statusBar()  #Jet：修改，添加状态栏
        font = QtGui.QFont()
        font.setPointSize(10)

        self.statusbar.setFont(font)
        self.statusbar.setObjectName("statusbar")
        self.setStatusBar(self.statusbar)
        # self.statusbar.showMessage('here is statusbar') # 在状态栏显示指定消息

        # self.retranslateUi(MainWindow)
        self.retranslateUi()

        # self.radioButton_s1iscoach.toggled['bool'].connect(self.radioButton_s1isclient.setChecked)

        QtCore.QMetaObject.connectSlotsByName(self)
        # print('here3.')

        # Jet: 按钮绑定相应事件,必须调用此函数！
        self.bind_events()

        # Jet: 添加以下这句，很关键，否则不显示窗口！
        self.show()

    # Jets：显示界面的文字。 从def retranslateUi(self, MainWindow)改写为:
    def retranslateUi(self):
        _translate = QtCore.QCoreApplication.translate
        self.setWindowTitle(_translate("MainWindow", "Jet's 音频转写-文字转换ICF格式"))

        # 第一行
        self.btn_locate_audio_file.setText(_translate("MainWindow", "1、请选择录音文件(mp3/wav)"))
        self.radioButton_1speaker.setText(_translate("MainWindow", "1位讲话者"))
        self.radioButton_2speakers.setText(_translate("MainWindow", "2位讲话者"))

        # 第二行
        self.btn_set_raw_output_file.setText(_translate("MainWindow", "2、设置原始文稿目录及文件名txt"))
        # btn提交讯飞原始转换
        self.btn_xunfei2txt.setText(_translate("MainWindow", "3、提交讯飞开始语音转换"))

        # 第三行
        self.radioButton_s2isclient.setText(_translate("MainWindow", "客户"))
        self.radioButton_s2iscoach.setText(_translate("MainWindow", "教练"))

        self.radioButton_s1iscoach.setText(_translate("MainWindow", "教练"))

        self.radioButton_s1isclient.setText(_translate("MainWindow", "客户"))

        self.label.setText(_translate("MainWindow", "请按实际情况指定：Speaker1："))
        self.label_2.setText(_translate("MainWindow", "Speaker2："))
        self.label_3.setText(_translate("MainWindow", "讯飞转写原始文稿："))
        self.label_4.setText(_translate("MainWindow", "ICF教练文稿格式："))
        # self.label_5.setText(_translate("MainWindow", "（注：讯飞转写可能存在少量错误，请根据实际录音编辑)"))

        self.btn_import_rawtxt.setText(_translate("MainWindow", "4、导入原始转写txt文件"))
        # 弃用  2020/07/09
        # self.btn_save_editing2stampfile.setText(_translate("MainWindow", "保存修改到时间戳文档"))
        # self.btn_import_stampbinary.setText(_translate("MainWindow", "导入编辑后的时间戳文档"))
        # self.btn_display_icf_format.setText(_translate("MainWindow", "5、导入原始数据为ICF格式"))

        # 添加按钮 2020/07/09
        self.btn_batch_edit.setText(_translate("MainWindow", "批量编辑原始文字"))
        self.ckbox_export_stat.setText(_translate("MainWindow", "包含统计信息"))

        self.btn_save_icf_to_binary.setText(_translate("MainWindow", "保存ICF格式到二进制文件"))
        self.btn_import_icf_from_binary.setText(_translate("MainWindow", "导入ICF格式二进制文件"))

        # self.btn_stop_audio.setText(_translate("MainWindow", "停止"))
        # self.btn_play_audio.setText(_translate("MainWindow", "播放"))

        self.ckbox_only_coach_text.setText(_translate("MainWindow", "只显示教练的文字"))

        self.btn_save_icf_to_txt.setText(_translate("MainWindow", "6、保存到txt文档"))
        self.btn_save_icf_to_word.setText(_translate("MainWindow", "7、保存到Word文档"))
        self.label_6.setText(_translate("MainWindow", "（注：讯飞转写可能存在少量错误，请根据实际录音编辑，满意之后再保存)"))

        # self.menu.setTitle(_translate("MainWindow", "教练录音转写助手"))

    # Jet: 用户选择用于转写文字的音频文件
    def btn_locate_audio_file_clicked(self):
        # show open file dialog
        self.cwd = os.getcwd()  # 获取当前程序文件位置

        fileName_choose, filetype = QFileDialog.getOpenFileName(self,
                                                                "选取文件",
                                                                self.cwd,  # 起始路径
                                                                "音频文件 (*.wav;*.mp3)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式

        if fileName_choose == "":
            # print("\n取消选择")
            # 警告：未选择音频文件
            QMessageBox.warning(self, '温馨提示', '您还没选择要转写的音频文件！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return
        # 修改当前路径为用户选择的音频文件目录
        # self.cwd = fileName_choose[0:-4] #不对，只是去年了后缀名！
        # 以上返回的是QString类型的对象，若想不出现编码问题，建议用如下语句将QString转换为python的string对象
        # fileName_choose = unicode(fileName_choose.toUtf8(), 'utf-8', 'ignore')

        # print("file:", fileName_choose)
        # 音频播放相关设置
        self.audio_file = fileName_choose  # 用于播放的音频文件名

        # 播放音质差，放弃使用
        # self.track = pygame.mixer.music.load(self.audio_file) #用pygame播放时需要此句

        # 按指定时间播放，差距太大，放弃使用
        self.mediaPlayer.setMedia(QMediaContent(QUrl.fromLocalFile(self.audio_file)))

        # self.audio_time = self.mediaPlayer.duration() #返回的是0，不能用！
        # self.audio_time = int(self.audio_time/ 1000)  #音频时长 seconds,返回的是0，不能用！
        # '''

        # 改用vlc播放
        self.vlc_player.set_uri(self.audio_file)
        self.vlcplayer_state_changed()

        self.audio_time = getaudioinfo.get_audio_time_seconds(self.audio_file)  # 音频时长 seconds
        self.btn_play_audio.setEnabled(True)  # 预备可播放音频

        self.slider_audio.setRange(0, self.audio_time * 1000)  # 决定当前视频的播放范围（ms）
        self.slider_audio.setMaximum(self.audio_time * 1000)
        self.slider_audio.setSingleStep(1000)
        self.slider_audio.setTickInterval(60000)
        self.slider_audio.setTickPosition(QSlider.TicksAbove)
        # '''
        # self.slider_audio.setRange(0, 100)
        self.cwd, tempfilename = os.path.split(fileName_choose)  # 分离文件路径和文件名
        self.filename_withoutextension, extension = os.path.splitext(tempfilename)  # 分离文件名和后缀
        '''
        import os
        file_path = "D:/test/test.py"
        (filepath, tempfilename) = os.path.split(file_path)
        (filename, extension) = os.path.splitext(tempfilename)
        '''

        # print("\n你选择的文件所在目录为:%s, 文件名为%s" % (self.cwd, self.filename_withoutextension))
        # print("文件筛选器类型: ",filetype)

        # 后续函数调用需用
        self.audio_file = fileName_choose

        # show the file path and name in the lineEdit
        self.lineEdit_audio_file.setText(fileName_choose)
        self.lineEdit_audio_file.setEnabled(False)

        # show default raw文件名
        self.raw_txt_file = fileName_choose[0:-4] + "_raw.txt"
        self.lineEdit_rawtxt_file.setText(self.raw_txt_file)

        self.statusbar.showMessage("您选择要转写的音频文件是" + fileName_choose)

        # 不提示：
        # QMessageBox.information(self, '下一步', '请选择第2步：设置保存原始转写文件的目录和文件名',
        #                        QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)

    # 用户选择保存原始转写结果的目录和文件名
    def btn_set_raw_output_file_clicked(self):
        # show save as dialog
        if self.cwd == "":
            self.cwd = os.getcwd()  # 获取当前程序文件位置
        # 否则使用和所选定音频文件所在的目录（见上一个函数）

        # self.filename_withoutextension
        fileName_choose, filetype = QFileDialog.getSaveFileName(self,
                                                                "文件保存为",
                                                                self.cwd,  # 起始路径
                                                                "Text Files (*.txt);;All Files (*)")  # 默认为txt格式
        # set default file name:

        if fileName_choose == "":
            # 警告：未选择文件名
            QMessageBox.warning(self, '温馨提示', '您还没选择要保存的文件名！', QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        # 后续函数调用需用
        self.raw_txt_file = fileName_choose  # 修改为用户输入的内容

        # show the file path and name in the lineEdit
        self.lineEdit_rawtxt_file.setText(fileName_choose)
        self.lineEdit_rawtxt_file.setEnabled(False)

        self.statusbar.showMessage("您选择要保存的原始转写的文件是" + fileName_choose)

        # Jet to add:...
        # save raw to file...

    # 调用Jet的讯飞转写函数，开始转写，显示转写进度的GUI, 存入用户指定的原始文件名
    def btn_xunfei2txt_clicked(self):

        if self.audio_file == "":
            # 警告：未选择文件名
            QMessageBox.warning(self, '温馨提示', '您还没选择要转写的音频文件！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        if self.raw_txt_file == "":
            # 警告：未选择文件名
            QMessageBox.warning(self, '温馨提示', '您还没选择要保存的文件名！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        self.statusbar.showMessage("开始转写音频文件" + self.audio_file)

        print("start call new Dialog:显示转写进度对话框，实时刷新转写情况")
        # Jet: 传递参数到子窗口,启动转写进度显示对话框
        # child = ChildWindow(window.main_ui.audio_file,
        #                    window.main_ui.raw_txt_file, window.main_ui.speaker_num)
        child = ConvertingDialog(self.audio_file, self.raw_txt_file, self.speaker_num)
        # 这3个参数由前两个函数调用时设置值
        child.exec_()  # 显示为model dialog模式对话框，未完成之前不得回到主窗口

        self.statusbar.showMessage("转写音频文件%s成功" % self.audio_file)

        # clear status bar msg

    # 从原始转写的raw文件导入，先转写成带时间戳的二进制文件，再显示到table widget中
    # 按教练/客户的对话显示（各人说话可能出现多次）
    # To modify: 从raw txt 判断是一人或二人的对话，再分别用不同的方式显示
    # 一人的说话，只需用起止时间戳，不用“讲话者”这一列信息
    def btn_import_rawtxt_clicked(self):
        """
        # modify on 2020/07/09
        调用display_icf_format， 显示ICF对话格式在右侧table widget中

        # Jet adds： 2020/06/30
        需要用到文件：raw_txt_file, 如果当前不存在，则需要选择文件

        :param raw_txt_file:
        :return:
        """
        self.raw_txt_file = self.lineEdit_rawtxt_file.text()  # 默认选择当前已经转写好的文件
        '''
        if not self.raw_txt_file:  # 如果当前不存在，则需要选择文件
            self.cwd = os.getcwd()  # 获取当前程序文件位置
            fileName_choose, filetype = QFileDialog.getOpenFileName(self,
                                                                    "选取文件",
                                                                    self.cwd,  # 起始路径
                                                                    "原始转写txt文件 (*.txt)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式
            if fileName_choose == "":
                # 警告：未选择文件
                QMessageBox.warning(self, '温馨提示', '您还没选择已经转写好的原始txt文件！',
                                    QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
                return
            self.raw_txt_file = fileName_choose
        '''
        # 2020-07-23 修改为如下：每次都需要指定转写好的raw文件
        self.cwd = os.getcwd()  # 获取当前凭什么目录
        fileName_choose, filetype = QFileDialog.getOpenFileName(self,
                                                                "选取文件",
                                                                self.cwd,  # 起始路径
                                                                "原始转写txt文件 (*.txt)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式
        if fileName_choose == "":
            # 警告：未选择文件
            QMessageBox.warning(self, '温馨提示', '您还没选择已经转写好的原始txt文件！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return
        self.raw_txt_file = fileName_choose

        self.cwd, tempfilename = os.path.split(self.raw_txt_file)  # 分离文件路径和文件名
        filename_withoutextension, extension = os.path.splitext(tempfilename)  # 分离文件名和后缀

        self.output_timestampbinary_file = filename_withoutextension + "_binary.data"

        # '''
        # 以下是按原始raw数据显示的，没经过整理

        if self.speaker_num == 2:

            # 确保每次显示的教练/客户是用户最新选择的
            if self.speaker1 == "教练":
                coach = "speaker1"
            else:
                coach = "speaker2"

            ''' 为了提高运行速度，不用先写入再读出
            # 调用转写函数-写入二进制文件中
            transform_tool.convert_raw2timestamp_dictlist_binary_file(self.raw_txt_file,
                                                                      self.output_timestampbinary_file,
                                                                      coach)
            # 从上述二进制文件中读取信息
            self.stamp_list = transform_tool.read_stamplist_from_binary_file(self.output_timestampbinary_file)
            '''
            # 2020-07-23 adds:
            self.icf_list = None  # 如果之前导入了raw文件，先清空.否则程序会卡
            self.icf_stat_list = None  # #如果之前导入了raw文件，先清空
            self.raw_output_result = None

            # 第一次调用就把从raw原始转换数据缓存到变量中，后续从中获取，以免重复调用函数
            self.raw_output_result, self.stamp_list = \
                transform_tool.get_stamp_list_from_rawtxt(self.raw_txt_file, coach)

            # 原始数据起止时间戳格式，显示到左侧table widget 中，需要先生成self.stamp_list
            self.display_2speakers_rawstamplist_in_table()

            # ICF对话格式,显示到右侧table widget 中, 包括显示统计信息到右下table widget中
            self.display_icf_format()

            # 显示统计信息到右下方的table widget中
            self.display_icfstat_in_table()

            # self.icf_stat_cache = self.icf_stat_list



        # 只有一位讲话者：
        elif self.speaker_num == 1:
            # to add...
            pass

    # 根据用户选择教练/用户的改变，只更新显示讲话者为客户/教练，并更新原始起止时间戳的内存
    def shiftspeakers_in_rawstamplist_table(self):
        if not self.stamp_list:
            return
        if self.speaker1 == "客户":  # 第一讲话者被改为客户，则交换缓存中的教练/客户
            for row, item in enumerate(self.stamp_list):
                # self.tableWidget_raw.setRowHeight(row, 30)
                # speak_time = item['speak_time']  # 07/02 updated in transform_tool.py
                speaker = item['speaker']  # 已经转写成了教练/客户
                # content = item['content']

                if speaker == "教练":
                    speaker = "客户"
                    item['speaker'] = "客户"
                else:
                    speaker = "教练"
                    item['speaker'] = "教练"

                # self.tableWidget_raw.setItem(row, 0, QTableWidgetItem(speak_time))
                self.tableWidget_raw.setItem(row, 1, QTableWidgetItem(speaker))  # 运行很慢，987条数据要花30-40秒！
                # 推测原因， tableWidget的刷新机制！

                # self.tableWidget_raw.setItem(row, 2, QTableWidgetItem(content))
                # row += 1
        print("finish iterate")

    # Jet adds on 2020/06/30
    # 显示到左侧table widget 中
    # def display_rawstamplist_in_table(self, stamp_list):
    # modified as follows on 2020/07/20 以便当用户更改了speaker1 = 教练/客户时，刷新显示左侧table
    def display_2speakers_rawstamplist_in_table(self):
        """
        需要先从raw.txt获取数据，调用函数是btn_import_rawtxt:
        self.raw_output_result, self.stamp_list = \
                transform_tool.get_stamp_list_from_rawtxt(self.raw_txt_file, coach)
        :return:
        """

        if not self.stamp_list:  # 空数据则不显示
            return
        # 设置table widget
        row_count = len(self.stamp_list)
        self.tableWidget_raw.clear()  # 清空

        self.tableWidget_raw.setColumnCount(3)
        self.tableWidget_raw.setRowCount(row_count)
        self.tableWidget_raw.setColumnWidth(0, 110)
        self.tableWidget_raw.setColumnWidth(1, 60)
        self.tableWidget_raw.setColumnWidth(2, 450)
        self.tableWidget_raw.horizontalHeader().setVisible(True)
        self.tableWidget_raw.verticalHeader().setVisible(True)

        # 设置水平表头
        # self.tableWidget_raw.setHorizontalHeaderLabels(["起止时间", "讲话者", "讲话内容"])
        header_items = [QTableWidgetItem("起止时间"), QTableWidgetItem("讲话者"), QTableWidgetItem("讲话内容")]
        for i, header_item in enumerate(header_items):
            self.set_header_item_format(header_item)  # 设置表头字体
            self.tableWidget_raw.setHorizontalHeaderItem(i, header_item)

        self.tableWidget_raw.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)  # 行高自适应
        # self.tableWidget_raw.resizeRowsToContents() #根据内容自动调整行高，没用！

        #  2020/07/03 设置为不可编辑
        # 2020/07/20 这样程序修改，更新显示时程序会死！故不用了。
        # self.tableWidget_raw.setEditTriggers(QAbstractItemView.NoEditTriggers)

        # '''第一次运行正常，第二次就很慢，死机（调试时没看到问题）
        row = 0
        for row, item in enumerate(self.stamp_list):
            '''
            第一次运行需要2秒，第二次运行需要29秒，奇怪！，比下面的方法要快
            -当切换了教练/客户之后，左下角table在刷新费时，有时要近一分钟！！！可能是table widget的问题
            ---解决办法：用分页显示（to do )
            '''
            self.tableWidget_raw.setRowHeight(row, 30)
            # start_time = item['start_time']
            speak_time = item['speak_time']  # 07/02 updated in transform_tool.py
            speaker = item['speaker']  # 已经转写成了教练/客户
            content = item['content']

            self.tableWidget_raw.setItem(row, 0, QTableWidgetItem(speak_time))
            self.tableWidget_raw.setItem(row, 1, QTableWidgetItem(speaker))
            self.tableWidget_raw.setItem(row, 2, QTableWidgetItem(content))
        # '''
        '''
        row = 0
        while row <  row_count:  # 980条数据需要30-40秒，花太长时间了！
            self.tableWidget_raw.setRowHeight(row, 30)
            item = self.stamp_list[row]
            # start_time = item['start_time']
            speak_time = item['speak_time']  # 07/02 updated in transform_tool.py
            speaker = item['speaker']  # 已经转写成了教练/客户
            content = item['content']

            self.tableWidget_raw.setItem(row, 0, QTableWidgetItem(speak_time))
            self.tableWidget_raw.setItem(row, 1, QTableWidgetItem(speaker))
            self.tableWidget_raw.setItem(row, 2, QTableWidgetItem(content))

            row += 1
        '''
        #print("finished.")
        # self.statusbar.showMessage("显示时间戳格式的原始数据")

    # 用户编辑完原始时间戳文件之后，保存到新的文件里（包括binary文件和txt文件）
    # 之后可导入此文件，合并同一位讲话者的内容，用ICF格式显示在右边的table widget里
    def btn_save_editing2stampfile_clicked(self):
        # 选择保存文件名

        # 如果还没有数据，则提示后退出
        row_count = self.tableWidget_raw.rowCount()
        if row_count == 0:
            QMessageBox.warning(self, '温馨提示', '您还没有可以存储的数据！', QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        fileName_choose, filetype = QFileDialog.getSaveFileName(self,
                                                                "文件保存为",
                                                                self.cwd,  # 起始路径
                                                                "时间戳二进制文件 (*.data);;All Files (*)")  # 默认为txt格式
        if fileName_choose == "":
            # 警告：未选择文件名
            QMessageBox.warning(self, '温馨提示', '您还没选择要保存的文件名！', QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        # 从table widget中读取最新编辑的数据，存到文件里（包括二进制文件和txt文件）
        dict_list = []
        # timestamp_dict = {}  #在这里会有Bug，存入有误！
        all_text = ""
        for i in range(row_count):
            start_time = self.tableWidget_raw.item(i, 0).text()
            speaker = self.tableWidget_raw.item(i, 1).text()
            content = self.tableWidget_raw.item(i, 2).text()

            """ 参考数据格式
            timestamp_dict = {
            "start_time":"00:01:21",
            "speaker": "speaker1",
            "content": "what the speaker says"
            }
            """
            timestamp_dict = {}
            timestamp_dict["start_time"] = start_time
            timestamp_dict["speaker"] = speaker
            timestamp_dict["content"] = content
            dict_list.append(timestamp_dict)

            new_item_text = start_time + " " + speaker + ": " + content + "\n"
            all_text = "".join([all_text, new_item_text])

        # 保存dict list到二进制文件，以便后续调用显示为ICF格式
        try:
            file_name = fileName_choose
            fw = open(file_name, 'wb')
            # print("dict_list:", str(dict_list))
            pickle.dump(dict_list, fw)
            # print('保存修改后的时间戳二进制文件:%s成功！' % file_name)
        except IOError:
            print('保存修改后的时间戳二进制文件:%s失败！' % file_name)
        finally:
            fw.close()

        # 保存文本到txt文件 (不必，可选，可供用户直接打开查看
        try:
            file_name = fileName_choose + ".txt"  # 默认的文件名
            fw = open(file_name, 'w', encoding='utf-8')
            fw.write(all_text)
            print('保存修改后的时间戳文件:%s成功！' % file_name)
        except IOError:
            print("保存失败!")
        finally:
            fw.close()

    # 播放/暂停声音
    def btn_play_audio_clicked(self):
        print("play audio:", self.audio_file)
        '''
        # 用pygame播放wav，mp3！
        if self.audio_file is not None:
            track = pygame.mixer.music.load(self.audio_file) #在选择文件时即初始化
            #To do : start_time = #从当前选择的table tablet中获取
            #pygame.mixer.music.play(start=start_time) #这里是按秒数？eg. 1.1
            print("playing audio now.")
            pygame.mixer.music.play()
        '''
        '''
        # 用QMediaPalyer播放
        if self.audio_file is not None:
            if self.mediaPlayer.state() == QMediaPlayer.PlayingState:
                self.mediaPlayer.pause()
            else:
                self.mediaPlayer.play()
        '''
        # 用vlc播放
        if self.audio_file is not None:
            if self.vlc_player.get_state() == 1:  # 如果在播放，则暂停
                self.vlc_player.pause()
            elif self.vlc_player.get_state() == 0:  # 如果已经暂停，则播放
                self.vlc_player.play()

        self.vlcplayer_state_changed()

    '''
    # 停止播放,不用pygame播放了
    def btn_stop_audio_clicked(self):
        if self.audio_file is not None:
            # track = pygame.mixer.music.load(self.audio_file) #还需要吗？
            pygame.mixer.music.stop()
            print("stop playing now.")
    '''

    # Jet adds on 2020/06/30
    # 显示到table widget 中
    # def display_stamplist_in_table(self, stamp_list):
    # 2020/07/27 modified above as follow:
    def display_stamplist_in_table(self):

        if not self.stamp_list or self.stamp_list is None:  # 空数据则不显示
            return
        # 设置table widget
        row_count = len(self.stamp_list)
        self.tableWidget_raw.clear()  # 清空

        self.tableWidget_raw.setColumnCount(3)
        self.tableWidget_raw.setRowCount(row_count)
        self.tableWidget_raw.setColumnWidth(0, 90)
        self.tableWidget_raw.setColumnWidth(1, 40)
        self.tableWidget_raw.setColumnWidth(2, 400)

        self.tableWidget_raw.horizontalHeader().setVisible(True)
        self.tableWidget_raw.verticalHeader().setVisible(True)

        # self.tableWidget_raw.setHorizontalHeaderLabels(["起止时间", "讲话者", "讲话内容"])
        header_items = [QTableWidgetItem("起止时间"), QTableWidgetItem("讲话者"), QTableWidgetItem("讲话内容")]
        for i, header_item in enumerate(header_items):
            # self.set_item_format(header_item)  # 设置表头字体
            self.set_header_item_format(header_item)  # 设置表头字体
            self.tableWidget_raw.setHorizontalHeaderItem(i, header_item)

        #
        # self.tableWidget_raw.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)  # 行高自适应
        # 行数多了会显示得慢，尤其是第二次刷新时！
        # 改为以下方式：没有变化！
        # self.tableWidget_raw.verticalHeader().setSectionResizeMode(QHeaderView.Fixed)
        # self.tableWidget_raw.horizontalHeader().setSectionResizeMode(QHeaderView.Fixed)

        # self.tableWidget_raw.resizeRowsToContents() #根据内容自动调整行高，没用！

        # 2020/07/03 设置为不可编辑
        # 2020/07/20 这样程序修改，更新显示时程序会死！故不用了。
        # self.tableWidget_raw.setEditTriggers(QAbstractItemView.NoEditTriggers)

        # row = 0
        for row, item in enumerate(self.stamp_list):
            self.tableWidget_raw.setRowHeight(row, 30)

            #start_time = item['start_time']
            start_time = item['speak_time']

            speaker = item['speaker']  # 这是已经被改写过的，内容： 教练/客户
            content = item['content']
            self.tableWidget_raw.setItem(row, 0, QTableWidgetItem(start_time))
            self.tableWidget_raw.setItem(row, 1, QTableWidgetItem(speaker))
            self.tableWidget_raw.setItem(row, 2, QTableWidgetItem(content))
            # row += 1

    # 从编辑后的时间戳二进制文件读取数据，显示到Table Widget中
    def btn_import_stampbinary_clicked(self):
        self.cwd = os.getcwd()  # 获取当前程序文件位置
        fileName_choose, filetype = QFileDialog.getOpenFileName(self,
                                                                "选取文件",
                                                                self.cwd,  # 起始路径
                                                                "时间戳文件 (*.data)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式
        if fileName_choose == "":
            # 警告：未选择文件
            QMessageBox.warning(self, '温馨提示', '您还没选择时间戳文件！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        self.stamp_list = transform_tool.read_stamplist_from_binary_file(fileName_choose)
        # print("reads stamp_list:", str(stamp_list))
        self.display_stamplist_in_table()

    # Jet adds on 2020/07/04
    # item = QTableWidgetItem('文字')
    # fore_color, bg_color: QBrush(QColor(255, 0, 0))
    #
    def set_item_format(self, item,
                        fore_color=QColor(0, 0, 0),
                        bg_color=QBrush(QColor(255, 255, 255)),
                        font_name="微软雅黑",
                        font_size=11):
        # 设置字体（字体，字号，颜色）
        item.setFont(QFont(font_name, font_size))
        item.setForeground(fore_color)
        # 另外一种方法设置颜色（该为红色）
        # item.setForeground(QBrush(QColor(255, 0, 0)))
        # 设置背景色
        item.setBackground(bg_color)

    # Jet adds on 2020/07/07
    # 设置表头格式
    def set_header_item_format(self, item,
                               fore_color=QColor(0, 0, 255),
                               bg_color=QBrush(QColor(255, 255, 255)),
                               font_name="微软雅黑",
                               font_size=11):
        # 设置字体（字体，字号，颜色）
        item.setFont(QFont(font_name, font_size, QFont.Bold))  # 设为粗体
        item.setForeground(fore_color)
        # 另外一种方法设置颜色（该为红色）
        # item.setForeground(QBrush(QColor(255, 0, 0)))
        # 设置背景色
        item.setBackground(bg_color)

    # Jet adds on 2020/07/03
    # 显示到icftable widget 中
    # def display_icflist_in_table(self, icf_list):
    # 2020/07/25 修改，删除参数，改用self.icf_list
    def display_icflist_in_table(self):

        if not self.icf_list:  # 空数据则不显示
            return
        # 设置table widget
        row_count = len(self.icf_list)
        self.tableWidget_icf.clear()  # 清空

        self.tableWidget_icf.setColumnCount(3)
        self.tableWidget_icf.setRowCount(row_count)
        self.tableWidget_icf.setColumnWidth(0, 110)
        self.tableWidget_icf.setColumnWidth(1, 60)
        self.tableWidget_icf.setColumnWidth(2, 450)
        self.tableWidget_icf.horizontalHeader().setVisible(True)
        self.tableWidget_icf.verticalHeader().setVisible(True)

        # 设置水平表头
        # self.tableWidget_icf.setHorizontalHeaderLabels(["起止时间", "讲话者", "讲话内容"])
        header_items = [QTableWidgetItem("起止时间"), QTableWidgetItem("讲话者"), QTableWidgetItem("讲话内容")]
        for i, header_item in enumerate(header_items):
            self.set_header_item_format(header_item)  # 设置表头字体
            self.tableWidget_icf.setHorizontalHeaderItem(i, header_item)

        # 行高自适应，特别是文字很多时
        self.tableWidget_icf.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        # Ref：https://blog.csdn.net/Heaven_Evil/article/details/78617461?utm_medium=distribute.pc_relevant.none-task-blog-BlogCommendFromMachineLearnPai2-3.nonecase&depth_1-utm_source=distribute.pc_relevant.none-task-blog-BlogCommendFromMachineLearnPai2-3.nonecase
        # self.tableWidget_icf.resizeRowsToContents()  # 根据内容自动调整行高，没用！

        # row = 0
        # for item in icf_list:
        # for row, dialog_item in enumerate(icf_list):  # 按编程规范修改：使用enumerate进行迭代，不用row = 0, row += 1。
        # 7/9 modified, 还得用row =0,否则会在只显示教练文字时，把客户文字显示为空白行
        row = 0
        for dialog_item in self.icf_list:
            speak_time = dialog_item['speak_time']  # 07/02 updated in transform_tool.py
            # 参见：convert_raw2icf_format_list函数
            speaker = dialog_item['speaker']
            content = dialog_item['content']

            widget_items = [QTableWidgetItem(speak_time), QTableWidgetItem(speaker), QTableWidgetItem(content)]

            # widget_item1 = QTableWidgetItem(speak_time)
            # widget_item2 = QTableWidgetItem(speaker)
            # widget_item3 = QTableWidgetItem(content)

            if self.is_only_coach_text and speaker == "客户":  # 当选择只显示教练的文字时，跳过客户的文字
                # row += 1  #不能有此句，否则会多出空白行
                continue
            for j, widget_item in enumerate(widget_items):
                if j == 0 or j == 1:
                    widget_item.setTextAlignment(QtCore.Qt.AlignHCenter)  # 第1，2列设为文字在单元格中居中
                # 设置教练的行数：前景色为蓝色，背景色为浅绿色（与客户的行不同）
                if speaker == "教练":
                    self.set_item_format(widget_item,
                                         fore_color=QBrush(QColor(0, 0, 250)),
                                         bg_color=QBrush(QColor(100, 250, 200)))
                    # 或一个一个地设置
                    # widget_item1.setFont(QFont("微软雅黑", 14))
                    # widget_item1.setForeground(QBrush(QColor(0, 0, 250)))
                    # widget_item1.setBackground(QBrush(QColor(102, 0, 68)))
                    # widget_items[j].setFont(QFont("微软雅黑", 14, QFont.Blue))
                    # 设置背景色
                    # widget_items[j].setBackground(QBrush(QColor(102, 0, 68)))
                else:
                    # widget_item1.setFont(QFont("微软雅黑", 14))
                    # widget_item1.setBackground(QBrush(QColor(255, 255, 255)))

                    self.set_item_format(widget_item)  # 客户的行，默认设置
                    # widget_items[j].setFont(QFont("微软雅黑", 14, QFont.Black))
                    # 设置背景色
                    # widget_items[j].setBackground(QBrush(QColor(255, 255, 255)))

                self.tableWidget_icf.setItem(row, j, widget_item)
            row += 1

    # 把ICF教练统计信息显示在下方table widget中
    # icf_stat_list 来自 convert_raw2icf_format_list函数返回的元组[1]
    # 2020/07/03 添加
    # def display_icfstat_in_table(self, icf_stat_list):
    # 2020/07/25 modifed, del para icf_stat_list, 改用 self.icf_stat_list
    def display_icfstat_in_table(self):

        if not self.icf_stat_list:  # 空数据则不显示
            return

        # 重新计算教练和客户讲话的字数，更新self.icf_stat_list
        self.icf_stat_list = self.calc_stat_from_icf_list()

        # 设置table widget

        self.tableWidget_stat.clear()  # 清空

        self.tableWidget_stat.setColumnCount(6)
        self.tableWidget_stat.setRowCount(2)
        self.tableWidget_stat.horizontalHeader().setVisible(True)
        self.tableWidget_stat.verticalHeader().setVisible(True)

        # 设置水平表头
        # self.tableWidget_stat.setHorizontalHeaderLabels(["说话者", "说话时长",
        #                                                 "时长占比", "安静时长", "字符总数", "字符占比"])

        header_items = [QTableWidgetItem("说话者"),
                        QTableWidgetItem("说话时长"),
                        QTableWidgetItem("时长占比"),
                        QTableWidgetItem("安静时长"),
                        QTableWidgetItem("字符总数"),
                        QTableWidgetItem("字符占比")]

        for i, header_item in enumerate(header_items):
            self.set_header_item_format(header_item)  # 设置表头字体

            self.tableWidget_stat.setHorizontalHeaderItem(i, header_item)

        # self.tableWidget_stat.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)  # 行高自适应
        # self.tableWidget_stat.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)  # 列宽自适应
        self.tableWidget_stat.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)  # 使表宽度自适应
        #  2020/07/07 设置为不可编辑
        self.tableWidget_stat.setEditTriggers(QAbstractItemView.NoEditTriggers)

        # row = 0
        # for speaker_dict in icf_stat_list:

        for row, speaker_dict in enumerate(self.icf_stat_list):
            widget_items = [QTableWidgetItem(speaker_dict["说话者"]),
                            QTableWidgetItem(speaker_dict["说话时长"]),
                            QTableWidgetItem(speaker_dict["时长占比"]),
                            QTableWidgetItem(speaker_dict["安静时长"]),
                            QTableWidgetItem(speaker_dict["字符总数"]),
                            QTableWidgetItem(speaker_dict["字符占比"])]

            for j, widget_item in enumerate(widget_items):
                self.set_item_format(widget_item)  # 设置表格默认设置
                self.tableWidget_stat.setItem(row, j, widget_item)
                widget_item.setTextAlignment(QtCore.Qt.AlignCenter)  # 居中

            # row += 1

    # 从raw导入ICF格式,并且显示在table view中，供用户查看、编辑
    # 显示统计信息到右下方的table widget中
    # def btn_display_icf_format_clicked(self):
    # 改名如下, 添加文件名参数，供btn_import_rawtxt_clicked调用 on 2020/07/09
    def display_icf_format(self):
        """
        根据原始raw转写，显示为ICF对话格式，同一个人的讲话放到一起

        区分是教练或客户（根据上方RadioButton的选择来确定教练/客户）
        --注意区分：raw: speaker = 1/2, Jet自定义  speaker = speaker1/speaker2
        加入安静的时间
        :return:
        """
        '''不用选择文件了
        raw_txt_file = self.lineEdit_rawtxt_file.text()  # 默认选择当前已经转写好的文件

        if not raw_txt_file:  # 如果当前不存在，则需要选择文件
            self.cwd = os.getcwd()  # 获取当前程序文件位置
            fileName_choose, filetype = QFileDialog.getOpenFileName(self,
                                                                    "选取文件",
                                                                    self.cwd,  # 起始路径
                                                                    "原始转写txt文件 (*.txt)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式
            if fileName_choose == "":
                # 警告：未选择文件
                QMessageBox.warning(self, '温馨提示', '您还没选择已经转写好的原始txt文件！',
                                    QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
                return
            raw_txt_file = fileName_choose

        self.cwd, tempfilename = os.path.split(raw_txt_file)  # 分离文件路径和文件名
        filename_withoutextension, extension = os.path.splitext(tempfilename)  # 分离文件名和后缀

        self.output_timestampbinary_file = filename_withoutextension + "_icflist_binary.data"
        '''
        if not self.raw_output_result or self.raw_output_result is None:
            return

        # 2020/07/03 改编： 按Jet整理过的数据格式来显示
        # who is coach? speaker1 or speaker2?
        if self.speaker1 == "教练":
            coach = "speaker1"
        else:
            coach = "speaker2"

        # 关键：从原始转写文件获取icf格式list和统计信息 stat_list
        # 2020/07/25 fix bug:
        # 第一次调用才从原始raw.txt获取，后续只需从缓存中调入，否则修改后的数据不会刷新！
        if not self.icf_list:
            self.icf_list, self.icf_stat_list, self.coaching_row_count = \
                transform_tool.convert_raw2icf_format_list(self.raw_output_result, coach)

        # 显示ICF对话模式到table widget 中
        #self.display_icflist_in_table(self.icf_list)
        self.display_icflist_in_table()

        # adds on 2020/07/09 update cache
        self.icf_list_cache = self.icf_list

    # 先用def get_icf_list_from_tablewidget(self)获取最新的icf_list
    # 计算最新的（用户可能修改过的）统计信息（esp.教练和客户说话的字数）
    # 在保存icf到txt/word之前需要调用！
    def calc_stat_from_icf_list(self):
        # 从表格获取最新icf数据（必须确保同时包含教练和客户的对话)
        if self.is_only_coach_text:
            QMessageBox.warning(self, '温馨提示', '您需要同时显示教练和客户的对话才能进行统计！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return
        else:
            icf_list = self.get_icf_list_from_tablewidget()
            """
            从self.tableWidget_icf中获取数据(用户可能编辑过的最新数据）
            item_dict = {
                "speak_time":"01:21-01:25",
                "speaker": "教练/客户",
                "content": "what the speaker says"
            }
            """
            # 分别统计（用户可能编辑之后的）教练和客户讲话的文字（中文）字数，忽略英文、数字及标点
            coach_text_count = 0
            client_text_count = 0
            for i, item in enumerate(icf_list):
                if item["speaker"] == "教练":
                    coach_text_count += transform_tool.count_zh_chars(item["content"])
                    if "【安静" in item["content"]:
                        coach_text_count -= 3  # 例如：减去 “【安静3秒】”中的三个汉字
                else:
                    client_text_count += transform_tool.count_zh_chars(item["content"])
                    if "【安静" in item["content"]:
                        client_text_count -= 3  # 减去 “【安静3秒】”中的三个汉字

            coach_words_ratio = coach_text_count / (coach_text_count + client_text_count)

            # print(" coach text zh count: ", coach_text_count)
            # print(" client text zh count: ", client_text_count)

            # 说话时长和安静时长需要从raw原始数据中获取！
            """
            self.icf_stat_list  # 这是导入raw txt时获得的,只调用以下函数一次
            参考transform_tool.convert_raw2icf_format_list(raw_txt_file, coach="speaker1"):
            函数从原始数据返回的统计信息：
            icf_stat_list:
                speaker1_dict["说话者"] = "教练"或"客户"
                speaker1_dict["说话时长"] = "t秒"
                speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
                speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
                speaker1_dict["字符总数"] = "%d" % total_length_speaker1
                speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)
            """

            # 注意：如果用户切换了speaker1为用户，则以上数据需要交换
            # to do : 当用户切换了speaker1的类型时，应当把self.icf_stat_list中的教练/客户交换

            # 以下默认数据是最新的了。
            # 用最新字数刷新字典的内容
            # to do : 以下逻辑需要理顺 2020-07-21 eve
            if self.icf_stat_list[0]["说话者"] == "教练":
                coach_stat_dict = self.icf_stat_list[0]
                self.icf_stat_list[0]["字符总数"] = "%d" % coach_text_count
                self.icf_stat_list[0]["字符占比"] = "%.2f%%" % (coach_words_ratio * 100)

                client_stat_dict = self.icf_stat_list[1]
                self.icf_stat_list[1]["字符总数"] = "%d" % client_text_count
                self.icf_stat_list[1]["字符占比"] = "%.2f%%" % ((1 - coach_words_ratio) * 100)

            elif self.icf_stat_list[0]["说话者"] == "客户":
                coach_stat_dict = self.icf_stat_list[1]
                self.icf_stat_list[1]["字符总数"] = "%d" % coach_text_count
                self.icf_stat_list[1]["字符占比"] = "%.2f%%" % (coach_words_ratio * 100)

                client_stat_dict = self.icf_stat_list[0]
                self.icf_stat_list[0]["字符总数"] = "%d" % client_text_count
                self.icf_stat_list[0]["字符占比"] = "%.2f%%" % ((1 - coach_words_ratio) * 100)

        return self.icf_stat_list

    # 当用户切换了speaker1为客户/教练之后，修改内存中icf_stat_list的教练/客户
    def switch_coach_client_in_stat(self):
        if self.icf_stat_list is None:
            return

        if self.speaker1 == "客户" and self.icf_stat_list[0]["说话者"] == "教练":
            self.icf_stat_list[0]["说话者"] == "客户"
            self.icf_stat_list[1]["说话者"] == "教练"
        elif self.speaker1 == "教练" and self.icf_stat_list[0]["说话者"] == "客户":
            self.icf_stat_list[0]["说话者"] == "教练"
            self.icf_stat_list[1]["说话者"] == "客户"

    # 2020/07/25 adds
    #当用户切换了教练/客户时，切换缓存icf_list教练/客户
    def switch_coach_client_in_icf_list(self):
        if self.icf_list is None:
            return

        for i, dialog_item in enumerate(self.icf_list):
            speaker = dialog_item['speaker']
            if speaker == "教练":
                self.icf_list[i]['speaker'] = "客户"
            else:
                self.icf_list[i]['speaker'] = "教练"

       # print(self.icf_list)

    # 2020/07/27 adds
    #当用户切换了教练/客户时，切换缓存stamp_list教练/客户,用于左下table显示的数据
    def switch_coach_client_in_stamp_list(self):
        if self.stamp_list is None:
            return
        for i, dialog_item in enumerate(self.stamp_list):
            speaker = dialog_item['speaker']
            if speaker == "教练":
                self.stamp_list[i]['speaker'] = "客户"
            else:
                self.stamp_list[i]['speaker'] = "教练"

    # 2020/07/04添加
    def get_icf_list_from_tablewidget(self):
        """
        从self.tableWidget_icf中获取数据(用户可能编辑过的最新数据）
        item_dict = {
            "speak_time":"01:21-01:25",
            "speaker": "教练/客户",
            "content": "what the speaker says"
        }

        :return:返回字典列表
        """
        row_count = self.tableWidget_icf.rowCount()
        if row_count == 0:
            # 提示，目前还没有ICF教练对话数据：
            QMessageBox.warning(self, '温馨提示', '您还没有教练对话数据！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return
        # 从self.tableWidget_icf中获取数据(用户可能编辑过）
        icf_list = []
        for i in range(row_count):
            # 以下两行代码很重要！！
            cell = self.tableWidget_icf.item(i, 1)  # 获取第二列：应当显示： 教练/客户
            if cell is not None and cell.text() != '':
                # 避免当只显示教练数据时，下方多余的单元格为空，程序出错
                item_dict = {"speak_time": self.tableWidget_icf.item(i, 0).text(),
                             "speaker": self.tableWidget_icf.item(i, 1).text(),
                             "content": self.tableWidget_icf.item(i, 2).text()}
                icf_list.append(item_dict)

        return icf_list

    # 2020/07/20添加
    def get_icf_coaching_list_from_tablewidget(self):
        """
        从self.tableWidget_icf中获取教练数据(用户可能编辑过的最新数据）
        item_dict = {
            "speak_time":"01:21-01:25",
            "speaker": "教练",
            "content": "what the speaker says"
        }

        :return:返回字典列表
        """
        row_count = self.coaching_row_count
        if row_count == 0:
            # 提示，目前还没有ICF教练对话数据：
            QMessageBox.warning(self, '温馨提示', '您还没有教练对话数据！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return
        # 从self.tableWidget_icf中获取数据(用户可能编辑过）
        icf_list = []
        for i in range(row_count):
            item_dict = {"speak_time": self.tableWidget_icf.item(i, 0).text(),
                         "speaker": self.tableWidget_icf.item(i, 1).text(),
                         "content": self.tableWidget_icf.item(i, 2).text()}
            icf_list.append(item_dict)

        return icf_list

    # 重新计算统计信息（在用户修改了ICF表格中的内容之后）
    # 前置：添加一个相应的按钮，对应的事件用本函数
    def recalc_icf_stat(self):
        icf_list = self.get_icf_list_from_tablewidget()
        # [安静]的信息保存在说话的文本中了，指定秒数以内若是安静，则这段时间被忽略了！
        # 只能大致计算，不能精确计算，总时长加起来会小于音频总时长！

        # To do ....

        # update display in stat table widget

    # 2020/07/04添加
    def get_icf_stat_from_tablewidget(self):
        """
        从self.tableWidget_stat中获取ICF教练统计数据
        :return:
        icf_stat_list:
            speaker1_dict["说话者"] = "教练"或"客户"
            speaker1_dict["说话时长"] = "t秒"
            speaker1_dict["时长占比"] = "%.2f%%" % (speaker1_time_ratio * 100)
            speaker1_dict["安静时长"] = "%d 秒" % speaker1_quiet_time
            speaker1_dict["字符总数"] = "%d" % total_length_speaker1
            speaker1_dict["字符占比"] = "%.2f%%" % (speaker1_words_ratio * 100)
        """

        # 从self.tableWidget_icf中获取数据(用户可能编辑过）
        icf_stat_list = []
        for i in range(self.tableWidget_stat.rowCount()):
            item_dict = {"说话者": self.tableWidget_stat.item(i, 0).text(),
                         "说话时长": self.tableWidget_stat.item(i, 1).text(),
                         "时长占比": self.tableWidget_stat.item(i, 2).text(),
                         "安静时长": self.tableWidget_stat.item(i, 3).text(),
                         "字符总数": self.tableWidget_stat.item(i, 4).text(),
                         "字符占比": self.tableWidget_stat.item(i, 5).text()}
            icf_stat_list.append(item_dict)

        return icf_stat_list

    # 音乐播放相关 --------------
    # """
    # 适用于QMediaPlayer
    def mediastate_changed(self, state):
        if self.mediaPlayer.state() == QMediaPlayer.PlayingState:
            self.btn_play_audio.setIcon(
                self.style().standardIcon(QStyle.SP_MediaPause)

            )

        else:
            self.btn_play_audio.setIcon(
                self.style().standardIcon(QStyle.SP_MediaPlay)

            )

    # """

    # 适用于vlc player，需要改写，在btn_play_audio_clicked函数中调用
    # 改变按钮图标， 不用这个办法，改用下面的VLC事件
    def vlcplayer_state_changed(self):
        pass
        '''
        if self.vlc_player.get_state() == 1:
            self.btn_play_audio.setIcon(
                self.style().standardIcon(QStyle.SP_MediaPause)

            )

        else:
            self.btn_play_audio.setIcon(
                self.style().standardIcon(QStyle.SP_MediaPlay)

            )
        '''

    # 针对vlc的回调函数，当播放时间改变时（VLC中的set_time被调用时），设置进度条
    def vlc_call_back_time_changed(self, event):
        audio_time = self.vlc_player.media.get_time()
        self.slider_audio.setValue(audio_time)

    # 更换播放按钮图标
    def vlc_call_back_playing(self, event):
        self.btn_play_audio.setIcon(
            self.style().standardIcon(QStyle.SP_MediaPause)

        )

    # 更换播放按钮图标
    def vlc_call_back_paused(self, event):
        self.btn_play_audio.setIcon(
            self.style().standardIcon(QStyle.SP_MediaPlay)

        )

    # 进度条调整时，或在Table Widget上选择的行对应的时间变化时触发
    def position_changed(self, position):
        self.slider_audio.setValue(position)
        # print("position:", self.mediaPlayer.position()) #for QMediaPlayer
        # print("position:", self.vlc_player.get_time())

    # 设置进度条已经播放的位置
    def duration_changed(self, duration):
        self.slider_audio.setRange(0, duration)

    # 设置播放毫秒数
    def set_position(self, ms):
        # self.mediaPlayer.setPosition(position)  # for QMediaPlayer
        # print("position:", position)

        # self.vlc_player.play()  # 必须先播放，
        # time.sleep(0.01)  # 稍微停一下,设置时间才管用
        # 以上两句已经加入到VLC_Jet.py中的set_time()中了
        self.vlc_player.set_time(ms)  # 从指定毫秒处播放

    # Jet adds: 根据TableWidge当前行的时间戳，设置起始播放的秒数
    # def set_play_start_seconds(self):
    # modified on 2020/07/09
    def set_play_start_seconds(self):
        # 如果选择了TableWidget中的一行（带有时间戳，则以此时间戳为准）

        # row_index = self.tableWidget_raw.currentIndex().row()  # 获取当前行Index
        # time_str = self.tableWidget_raw.item(row_index, 0).text()  # item(行,列), 获取当前行第一例,起始时间

        tableWidget = self.sender()  # modified on 2020/07/09
        # 可能是tableWidget_raw 或tableWidget_icf
        row_index = tableWidget.currentIndex().row()  # 获取当前行Index
        time_str = tableWidget.item(row_index, 0).text()  # item(行,列), 获取当前行第一例,起始时间

        # 2020/07/03 从开始时间改为了起止时间00:01-00:03，修改程序如下：
        time_list = time_str.split("-")
        start_time_str = time_list[0]

        str_list = start_time_str.split(":")
        m = int(str_list[0])
        s = int(str_list[1])

        self.start_seconds = m * 60 + s
        # ms = self.start_seconds *1000/2
        ms = self.start_seconds * 1000
        # print("set position in ms:", ms)
        # self.mediaPlayer.setPosition(ms)  #决定当前视频的播放位置（ms）# 为了更精准，还需要从原始数据获取毫秒数
        self.set_position(ms)  # 调用上面定义的此函数，就会触发按新位置播放

    # ------------------

    # adds on 2020/07/09
    # 从默认ICF icf_binary.data获取icf_list
    def import_icf_list_from_binary(self):
        if self.cwd == "":
            self.cwd = os.getcwd()  # 获取当前程序文件位置
        file_name = self.cwd + r"/icf_binary.data"
        try:
            fr = open(file_name, 'rb')  # b代表用二进制方式读取文件
            icf_list = pickle.load(fr)
            print("读取二进制文件%s成功！" % file_name)
        except IOError:
            print("读取二进制文件%s失败！" % file_name)
        finally:
            fr.close()

        return icf_list

    # adds on 2020/07/09
    # 把获取的icf_list保存到 icf_binary.data
    def save_icf_list_to_binary(self, icf_list):
        if self.cwd == "":
            self.cwd = os.getcwd()  # 获取当前程序文件位置
        file_name = self.cwd + r"/icf_binary.data"
        try:
            fw = open(file_name, 'wb')
            pickle.dump(icf_list, fw)
            print('保存修改后的时间戳二进制文件:%s成功！' % file_name)
        except IOError:
            print('保存修改后的时间戳二进制文件:%s失败！' % file_name)
        finally:
            fw.close()

    # adds on 2020/07/04
    # modified on 07/09, add parameter choose_File
    # def btn_save_icf_to_binary_clicked():
    def btn_save_icf_to_binary_clicked(self):
        """
        把当前icf 表格中的内容保存列表对象到二进制文件
        :return:
        """
        # 如果还没有数据，则提示后退出
        row_count = self.tableWidget_icf.rowCount()
        if row_count == 0:
            QMessageBox.warning(self, '温馨提示', '您还没有可以存储的数据！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        icf_list = self.get_icf_list_from_tablewidget()  # 从表格获取icf格式对话数据

        if self.cwd == "":
            self.cwd = os.getcwd()  # 获取当前程序文件位置
        # 指定保存的文件路径和文件名：
        file_name, filetype = QFileDialog.getSaveFileName(self, "文件保存为",
                                                          self.cwd,  # 起始路径
                                                          "二进制数据文件 (*.data);;All Files (*)")
        if file_name == "":
            # 警告：未选择文件名
            QMessageBox.warning(self, '温馨提示', '您还没选择要保存的文件名！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return
        # 保存dict list到二进制文件，以便后续调用显示为ICF格式
        try:
            fw = open(file_name, 'wb')
            pickle.dump(icf_list, fw)
            print('保存修改后的时间戳二进制文件:%s成功！' % file_name)
        except IOError:
            print('保存修改后的时间戳二进制文件:%s失败！' % file_name)
        finally:
            fw.close()

    def btn_import_icf_from_binary_clicked(self):
        """
        从指定(先前保存好的）icf格式对象二进制文件读取icf_list, 并且显示到icf table widget中
        :return:
        """
        # 选择之前已经保存好的二进制data文件
        self.cwd = os.getcwd()  # 获取当前程序文件位置

        file_name, filetype = QFileDialog.getOpenFileName(self,
                                                          "选取文件",
                                                          self.cwd,  # 起始路径
                                                          "ICF格式二进制数据文件 (*.data)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式
        if file_name == "":
            # 警告：未选择文件
            QMessageBox.warning(self, '温馨提示', '您还没选择二进制数据文件！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return

        try:
            fr = open(file_name, 'rb')  # b代表用二进制方式读取文件
            self.icf_list = pickle.load(fr)
        except IOError:
            print("读取二进制文件%s失败！" % file_name)
        finally:
            fr.close()

        # 数据有效性初步检测
        # if icf_list is None or len(icf_list) == 0:

        if not self.icf_list:
            return
        # to add： 数据格式检测(因为可能导入不同格式的data文件）
        # ...

        # 如果没有统计信息，则需要选择原始转写文件，计算出统计信息，并且保存到缓存里
        if not self.icf_stat_cache:
            self.cwd = os.getcwd()  # 获取当前程序文件位置

            raw_txt_file, filetype = QFileDialog.getOpenFileName(self,
                                                                 "请选取原始转写txt文件",
                                                                 self.cwd,  # 起始路径
                                                                 "原始转写txt文件 (*.txt)")  # 设置文件扩展名过滤,用双分号间隔;目前只支持这两种音频格式
            if raw_txt_file == "":
                # 警告：未选择文件
                QMessageBox.warning(self, '温馨提示', '您还没选择原始转写txt文件！',
                                    QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
                return

            if self.speaker_num == 2:
                if self.speaker1 == "教练":
                    coach = "speaker1"
                else:
                    coach = "speaker2"
                # icf_list0没用，不需要
                icf_list0, self.icf_stat_list = \
                    transform_tool.convert_raw2icf_format_list(self.raw_output_result, coach)

                # 把原始文件存到临时文件，再读出，显示到左侧table widget中
                # -----------
                self.cwd, tempfilename = os.path.split(raw_txt_file)  # 分离文件路径和文件名
                filename_withoutextension, extension = os.path.splitext(tempfilename)  # 分离文件名和后缀

                self.output_timestampbinary_file = filename_withoutextension + "_binary.data"

                transform_tool.convert_raw2timestamp_dictlist_binary_file(raw_txt_file,
                                                                          self.output_timestampbinary_file,
                                                                          coach)
                # 从上述二进制文件中读取信息
                self.stamp_dict_list = transform_tool.read_stamplist_from_binary_file(
                    self.output_timestampbinary_file)

                # 显示到左侧table widget 中
                self.display_2speakers_rawstamplist_in_table(self.stamp_dict_list)

                # -----------

                # 显示到icf table widget中
                #self.display_icflist_in_table(self.icf_list)
                self.display_icflist_in_table()

                # adds on 2020/07/09 update cache
                self.icf_list_cache = self.icf_list
                self.icf_stat_cache = self.icf_stat_list

                # 显示原始统计信息（或按编辑更新统计后的信息）到下方的控件中
                # to add...
                # 目前先用原始统计信息
                self.ckbox_export_stat.setChecked(True)
                # 上句不会触发调用下面这一句
                self.display_icfstat_in_table(self.icf_stat_list)

        '''
        参考：
        self.icf_list, self.icf_stat_list = transform2timestamp_txt_word.convert_raw2icf_format_list(raw_txt_file, coach)
        # 显示ICF对话模式到table widget 中
        self.display_icflist_in_table(self.icf_list)
        
        # 显示统计信息到下方的控件中
        self.display_icfstat_in_table(self.icf_stat_list)

        '''

    # Jet adds on 2020/07/07
    # 复选框，是否输出统计信息
    def chbox_state(self, ckbox):
        # self.ckbox_export_stat.stateChanged()
        ''' 测试：
        chk1Status = self.ckbox_export_stat.text() + ", isChecked=" + str(self.ckbox_export_stat.isChecked()) +\
                     ', checkState=' + str(self.ckbox_export_stat.checkState()) + "\n"

        print(chk1Status)
        '''
        # 1、是否显示统计内容
        self.is_stat_included = self.ckbox_export_stat.isChecked()

        if self.is_stat_included:
            # 从暂存二进制文件读出来，更新显示到icf table中
            # if not self.icf_list_cache:
            #   self.icf_list_cache = self.import_icf_list_from_binary()
            # else: 直接从缓存读取
            self.display_icfstat_in_table(self.icf_stat_cache)
        else:
            self.tableWidget_stat.clear()  # 不显示
            # 缓存数据到内存和文件中
            # self.icf_list_cache = self.get_icf_list_from_tablewidget()
            # self.save_icf_list_to_binary(self.icf_list_cache)

        # 2、是否只显示教练的文字
        self.is_only_coach_text = self.ckbox_only_coach_text.isChecked()

        # update UI display in icf table
        if self.is_only_coach_text:  # 先暂存数据到二进制文件里(和缓存里）,以便保留用户最新的编辑内容
            self.icf_list_cache = self.get_icf_list_from_tablewidget()
            self.save_icf_list_to_binary(self.icf_list_cache)
        else:
            # 从暂存二进制文件读出来，更新显示到icf table中
            if not self.icf_list_cache:
                self.icf_list_cache = self.import_icf_list_from_binary()
            else:
                pass  # 从缓存里读取之前保存到self.icf_list_cache的内容
            # 暂存教练和客户对话的数据，以便在选中only coach text, 再取消之后，用此数据恢复显示
        #self.display_icflist_in_table(self.icf_list_cache)
        self.display_icflist_in_table()

    # Jet adds on 2020/07/08
    def check_data_availability(self, file_format="txt"):
        """
        1、判断是否有可用数据
        2、选择文件保存路径
        file_format: txt  or word(docx)

        :return:  icf_list, stat_list, file_name
        """
        # 如果还没有数据，则提示后退出
        row_count = self.tableWidget_icf.rowCount()
        if row_count == 0:
            QMessageBox.warning(self, '温馨提示', '您还没有可以存储的ICF对话数据！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return False

        icf_list = self.get_icf_list_from_tablewidget()  # 从表格获取icf格式对话数据

        stat_list = self.get_icf_stat_from_tablewidget()  # 获取统计信息

        if not stat_list:
            QMessageBox.warning(self, '温馨提示', '您还没有可以存储的统计数据！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return False

        # 指定保存的文件路径和文件名：
        if self.cwd == "":
            self.cwd = os.getcwd()  # 获取当前程序文件位置

        if file_format == "txt":
            filetype = "文本文件 (*.txt);;All Files (*)"
        elif file_format == "docx":
            filetype = "Word文件 (*.docx);;All Files (*)"

        file_default_name = self.cwd + self.raw_txt_file[0:-4]  # # 起始路径 + 默认已经转写好的文本文件名（去掉后缀名.txt）
        file_name, filetype = QFileDialog.getSaveFileName(self, "文件保存为", file_default_name, filetype)
        if file_name == "":
            # 警告：未选择文件名
            QMessageBox.warning(self, '温馨提示', '您还没选择要保存的文件名！',
                                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            return False
        return icf_list, stat_list, file_name

    # Jet adds on 2020/07/07
    # 保存ICF对话模式到指定的txt文档,测试成功！
    def btn_save_icf_to_txt_clicked(self):
        """
        1、判断是否输出统计信息
        2、保存到指定txt文件

        :return: 保存状态True/False
        """
        results = self.check_data_availability(file_format="txt")
        icf_list, stat_list, file_name = results[0], results[1], results[2]

        str_time = icf_list[len(icf_list) - 1]["speak_time"]  # 最后一个记录的起止时间，eg: 59:37-59:39
        self.audio_time = str_time.split('-')[1]

        # 保存dict list到文本文件，以便后续调用显示为ICF格式
        try:
            fw = open(file_name, 'w', encoding='utf-8')
            text = ''
            # 是否输出统计信息：
            if self.is_stat_included and stat_list:  # 确保已经有统计数据
                # 表头
                for key in stat_list[0]:
                    text = ''.join([text, key, '\t'])  # Tab分隔
                text += "\n"  # 换行
                # '''
                for item_dict in stat_list:  # 分别打印教练和客户的统计信息
                    for key in item_dict:
                        text = ''.join([text, item_dict[key], '\t'])  # Tab分隔

                    text += "\n"  # 换行
            # print("stat:", text)

            text = ''.join([text, "教练时长：%s\n" % self.audio_time])

            # 教练对话转换为文字：
            text = ''.join([text, "教练对话笔录：\n"])

            for item_dict in icf_list:  # 分别打印教练和客户的统计信息

                # for key in item_dict:
                #    text = ''.join([text, item_dict[key], '\t'])  # Tab分隔太宽了
                if not self.is_only_coach_text:  # 分别打印教练和客户的统计信息
                    text = ''.join([text,
                                    item_dict["speak_time"],
                                    '[', item_dict["speaker"], ']',
                                    item_dict["content"]])
                    text += "\n"  # 换行
                elif item_dict["speaker"] == "教练":  # 只打印教练说的文字
                    text = ''.join([text,
                                    item_dict["speak_time"],
                                    '[', item_dict["speaker"], ']',
                                    item_dict["content"]])
                    text += "\n"  # 换行
            fw.write(text)
            print('保存到文件:%s成功！' % file_name)
        except IOError:
            print('保存到文件:%s失败！' % file_name)
        finally:
            fw.close()

    # modified on 2020-07-24
    def btn_save_icf_to_word_clicked(self):
        """
        Jet adds on 2020/07/24
        保存为ICF对话格式，根据用户选择，可添加统计数据到文件一开始
        1、判断是否输出统计信息
        2、保存到指定docx文件

        :return: 保存状态True/False
        """
        # 从现有表获取数据，设置另存为的文件名
        # 检验是否已经有数据
        results = self.check_data_availability(file_format="docx")
        if results:
            # icf_list, icf_stat_list, file_name = results[0], results[1], results[2]
            self.icf_list, self.icf_stat_list, file_name = results[0], results[1], results[2]

            # update stat data 按照最新编辑的icf对话文字，刷新统计数字
            self.icf_stat_list = self.calc_stat_from_icf_list()

            if transform_tool.save_icf_list_to_word(self.icf_list, self.icf_stat_list, file_name):
                QMessageBox.information(self, '提示', 'ICF对话已经成功保存到文档%s中！' % file_name, QMessageBox.Ok)
                # 直接打开此word文档. adds on 2020/10/08
                os.startfile(file_name)

            else:
                QMessageBox.warning(self, '提示', 'ICF对话保存到文档%s中失败！' % file_name, QMessageBox.Ok)

    # 用上面的函数简化本函数
    def btn_save_icf_to_word_clicked_bak(self):
        """
        Jet adds on 2020/07/08
        保存为ICF对话格式，根据用户选择，可添加统计数据到文件一开始
        1、判断是否输出统计信息
        2、保存到指定docx文件

        :return: 保存状态True/False
        """
        # 从现有表获取数据，设置另存为的文件名
        results = self.check_data_availability(file_format="docx")
        icf_list, stat_list, file_name = results[0], results[1], results[2]

        # ----------------
        try:
            # 打开文件
            document = docx.Document()

            # 以下会导致程序崩溃！
            # document = docx.Document("ICF-Word-template.docx")  #用Jet自定义的模板（带页眉和页脚，有页码）
            # '''添加页码，以下代码OK，注意需要import OxmlElement和os
            def create_element(name):
                return OxmlElement(name)

            def create_attribute(element, name, value):
                # element.set(docx.oxml.ns.qn(name), value)
                element.set(ns.qn(name), value)

            def add_page_number(paragraph):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                page_run = paragraph.add_run()
                t1 = create_element('w:t')
                create_attribute(t1, 'xml:space', 'preserve')
                t1.text = '第 '
                page_run._r.append(t1)

                page_num_run = paragraph.add_run()

                fldChar1 = create_element('w:fldChar')
                create_attribute(fldChar1, 'w:fldCharType', 'begin')

                instrText = create_element('w:instrText')
                create_attribute(instrText, 'xml:space', 'preserve')
                instrText.text = "PAGE"

                fldChar2 = create_element('w:fldChar')
                create_attribute(fldChar2, 'w:fldCharType', 'end')

                page_num_run._r.append(fldChar1)
                page_num_run._r.append(instrText)
                page_num_run._r.append(fldChar2)

                of_run = paragraph.add_run()
                t2 = create_element('w:t')
                create_attribute(t2, 'xml:space', 'preserve')
                # t2.text = '//' # Jet: 用转义反斜杠，隔开页码， 会出现//
                t2.text = r'/'  # Jet: 用转义反斜杠，隔开页码

                of_run._r.append(t2)

                fldChar3 = create_element('w:fldChar')
                create_attribute(fldChar3, 'w:fldCharType', 'begin')

                instrText2 = create_element('w:instrText')
                create_attribute(instrText2, 'xml:space', 'preserve')
                instrText2.text = "NUMPAGES"

                fldChar4 = create_element('w:fldChar')
                create_attribute(fldChar4, 'w:fldCharType', 'end')

                num_pages_run = paragraph.add_run()
                num_pages_run._r.append(fldChar3)
                num_pages_run._r.append(instrText2)
                num_pages_run._r.append(fldChar4)

                page_run = paragraph.add_run()
                t1 = create_element('w:t')
                create_attribute(t1, 'xml:space', 'preserve')
                t1.text = ' 页'
                page_run._r.append(t1)

            # 在页脚添加页码
            add_page_number(document.sections[0].footer.paragraphs[0])
            document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.save(file_name)
            # '''

            ''' 以下代码，在页脚添加页码，测试成功
            def create_element(name):
                return OxmlElement(name)

            def create_attribute(element, name, value):
                element.set(ns.qn(name), value)

            def add_page_number(run):
                fldChar1 = create_element('w:fldChar')
                create_attribute(fldChar1, 'w:fldCharType', 'begin')

                instrText = create_element('w:instrText')
                create_attribute(instrText, 'xml:space', 'preserve')
                instrText.text = "PAGE"

                fldChar2 = create_element('w:fldChar')
                create_attribute(fldChar2, 'w:fldCharType', 'end')

                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)

            add_page_number(document.sections[0].footer.paragraphs[0].add_run())
            
            
            document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.save(file_name)
            '''

            # 添加指定格式的段落/文字
            # text：段落内容
            def add_formatted_text(document, text, font_cn_name="宋体", size=12,
                                   color=RGBColor(0, 0, 0), alignment="left",
                                   is_bold=False):
                # 添加段落
                p1 = document.add_paragraph()
                text1 = p1.add_run(text)

                text1.font.name = 'Times New Roman'  # 控制是西文时的字体
                text1.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn_name)  # 控制是中文时的字体

                text1.font.size = Pt(size)
                text1.font.color.rgb = color

                # 此处高亮显示“安静”为绿色， 如果有的话
                # 此处还需改进。。。教练的文字，字体不对，颜色蓝色也没有了
                # highlight_keywords(p1, text, keyword, color)

                # 加粗
                text1.font.bold = is_bold
                # 斜体
                # text1.font.italic = True

                if alignment == "center":
                    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment == "left":
                    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                p1.paragraph_format.space_after = Pt(10)  # 段后间距，default

                document.save(file_name)

            # 第一段（设置标题居中）
            title = '教练对话语音转写'
            add_formatted_text(document, title, font_cn_name="宋体", size=16,
                               color=RGBColor(0, 0, 255), alignment="center",
                               is_bold=True)

            str_time = icf_list[len(icf_list) - 1]["speak_time"]  # 最后一个记录的起止时间，eg: 59:37-59:39
            self.audio_time = str_time.split('-')[1]
            total_time = '教练时长：%s' % self.audio_time
            add_formatted_text(document, total_time)

            # 是否输出统计信息：
            if self.is_stat_included and self.icf_list_cache:  # 确保已经有统计数据
                # 添加表格, 显示统计教练和客户的信息
                table = document.add_table(rows=3, cols=6,
                                           style="Light List Accent 5")
                # 指定表格格式，参见：https://blog.csdn.net/ibiao/article/details/78595295

                table.style.font.name = '宋体'  # 表格内的字体类型
                table.style.font.size = 120000  # 字号

                # 在表的指定行，列，输入内容
                def add_text2table(table, row, col, text):
                    cell = table.cell(row, col)
                    cell.text = text

                # 表格各列标题
                row, col = 0, 0
                for key in stat_list[0]:
                    add_text2table(table, row, col, key)
                    col += 1

                # 分别打印教练和客户的统计信息
                row, col = 1, 0
                for value in stat_list[0].values():  # 注意：必须加上.values(),否则只返回keys
                    add_text2table(table, row, col, value)
                    col += 1

                row, col = 2, 0
                for value in stat_list[1].values():
                    add_text2table(table, row, col, value)
                    col += 1

            text = "教练对话笔录:"
            add_formatted_text(document, text)

            '''
            # 只打印教练讲话的文字
            if self.is_only_coach_text:
                icf_coaching_list = self.get_icf_coaching_list_from_tablewidget()
                for item_dict in icf_coaching_list:  # 分别打印教练和客户的统计信息
                    #设置格式为：
                    #00:01-00:02[教练]喂，谭*你好！
                    text = ''.join([item_dict["speak_time"], ' [', item_dict["speaker"], '] ', item_dict["content"]])
                    # 教练讲话的字体设为蓝色：
                    if item_dict["speaker"] == "教练":
                        add_formatted_text(document, text, color=RGBColor(0, 0, 255))
            # 打印教练和客户对话数据
            else:
                for item_dict in icf_list:  # 分别打印教练和客户的统计信息
                    #设置格式为：
                    #00:01-00:02[教练]喂，谭*你好！
                    #0:03-00:05[客户] Hello。李**
                    text = ''.join([item_dict["speak_time"], ' [', item_dict["speaker"], '] ', item_dict["content"]])
                    #教练讲话的字体设为蓝色：
                    if item_dict["speaker"] == "教练":
                        add_formatted_text(document, text, color=RGBColor(0, 0, 255))
                    #elif item_dict[
                    #    "speaker"] == "客户" and not self.ckbox_only_coach_text:  # 客户讲话的字体用黑色,如果选中只输出教练文字时，则不输出客户文字
                    # 07/09 fix above bug by adding .isChecked()
                    elif item_dict["speaker"] == "客户" and not self.ckbox_only_coach_text.isChecked(): # 客户讲话的字体用黑色,如果选中只输出教练文字时，则不输出客户文字
                        add_formatted_text(document, text)
                        
            '''

            # 用户在右侧icf_table中看到什么，就打印什么（只有教练的讲话或包括教练和客户的对话）
            for item_dict in icf_list:  # 分别打印教练和客户的统计信息
                '''
                    设置格式为：
                    00:01-00:02[教练]喂，谭*你好！
                    00:03-00:05[客户] Hello。李**
                    '''
                text = ''.join([item_dict["speak_time"], ' [', item_dict["speaker"], '] ', item_dict["content"]])
                # 教练讲话的字体设为蓝色：
                if item_dict["speaker"] == "教练":
                    add_formatted_text(document, text, color=RGBColor(0, 0, 255))
                # elif item_dict[
                #    "speaker"] == "客户" and not self.ckbox_only_coach_text:  # 客户讲话的字体用黑色,如果选中只输出教练文字时，则不输出客户文字
                # 07/09 fix above bug by adding .isChecked()
                # 客户讲话的字体用黑色,如果选中只输出教练文字时，则不输出客户文字
                elif item_dict["speaker"] == "客户" and not self.ckbox_only_coach_text.isChecked():
                    add_formatted_text(document, text)

            print("保存结果到{}成功。".format(file_name))
        except IOError:
            print("保存结果到{}失败。".format(file_name))
        # ----------------

    # 添加 on 2020/07/09
    # 用于批量修改讯飞转写中的常见错误
    def btn_batch_edit_clicked(self):
        # 1、获取 ICF table数据
        '''
        从self.tableWidget_icf中获取数据
        item_dict = {
            "speak_time":"01:21-01:25",
            "speaker": "教练/客户",
            "content": "what the speaker says"
        }
        :return:
        '''
        self.icf_list = self.get_icf_list_from_tablewidget()
        if not self.icf_list:
            return

        # icf_list_new = [] # 存储批量修改之后的数据

        # 2、批量编辑，只需修改对话内容
        for dialog_item in self.icf_list:
            content = dialog_item['content']
            dialog_item['content'] = transform_tool.edit_common_error(content)

        # 3、刷新显示到ICF table中
        #self.display_icflist_in_table(self.icf_list)
        self.display_icflist_in_table()

        # 更新内存stat_list数据
        self.icf_stat_list = self.calc_stat_from_icf_list()

        return self.icf_list

    # 绑定按钮事件，
    # 需要被函数initUI调用
    def bind_events(self):

        self.btn_locate_audio_file.clicked.connect(self.btn_locate_audio_file_clicked)
        self.radioButton_1speaker.toggled.connect(self.onRadioButton_SpeakerNum_Clicked)
        self.radioButton_2speakers.toggled.connect(self.onRadioButton_SpeakerNum_Clicked)

        self.btn_set_raw_output_file.clicked.connect(self.btn_set_raw_output_file_clicked)
        self.btn_xunfei2txt.clicked.connect(self.btn_xunfei2txt_clicked)
        self.btn_import_rawtxt.clicked.connect(self.btn_import_rawtxt_clicked)

        # 弃用  2020/07/09
        # self.btn_save_editing2stampfile.clicked.connect(self.btn_save_editing2stampfile_clicked)
        # self.btn_import_stampbinary.clicked.connect(self.btn_import_stampbinary_clicked)
        # self.btn_display_icf_format.clicked.connect(self.btn_display_icf_format_clicked)
        # 添加按钮 2020/07/09
        self.btn_batch_edit.clicked.connect(self.btn_batch_edit_clicked)

        self.btn_save_icf_to_binary.clicked.connect(self.btn_save_icf_to_binary_clicked)
        self.btn_import_icf_from_binary.clicked.connect(self.btn_import_icf_from_binary_clicked)

        self.btn_play_audio.clicked.connect(self.btn_play_audio_clicked)
        # self.btn_stop_audio.clicked.connect(self.btn_stop_audio_clicked)
        self.radioButton_s1iscoach.toggled.connect(self.onRadioButtonClicked)  # Jet:注意，函数调用不能加()
        self.radioButton_s1isclient.toggled.connect(self.onRadioButtonClicked)
        self.radioButton_s2iscoach.toggled.connect(self.onRadioButtonClicked2)
        self.radioButton_s2isclient.toggled.connect(self.onRadioButtonClicked2)

        self.ckbox_export_stat.stateChanged.connect(lambda: self.chbox_state(self.ckbox_export_stat))

        self.ckbox_only_coach_text.stateChanged.connect(lambda: self.chbox_state(self.ckbox_only_coach_text))

        self.slider_audio.sliderMoved.connect(self.set_position)
        # 选择某行，就按此行开始时间播放语音
        self.tableWidget_raw.clicked.connect(self.set_play_start_seconds)
        # added on 2020/07/09
        # 选择某行，就按此行开始时间播放语音
        self.tableWidget_icf.clicked.connect(self.set_play_start_seconds)

        self.btn_save_icf_to_txt.clicked.connect(self.btn_save_icf_to_txt_clicked)
        self.btn_save_icf_to_word.clicked.connect(self.btn_save_icf_to_word_clicked)

        '''
        # 以下三个适用于QMediayPlayer(放弃使用，改用VLC）
        self.mediaPlayer.stateChanged.connect(self.mediastate_changed)
        self.mediaPlayer.positionChanged.connect(self.position_changed)
        self.mediaPlayer.durationChanged.connect(self.duration_changed)
        '''

        # 添加VLC事件监听（事件类型，槽函数）
        self.vlc_player.add_callback(vlc.EventType.MediaPlayerTimeChanged, self.vlc_call_back_time_changed)
        self.vlc_player.add_callback(vlc.EventType.MediaPlayerPlaying, self.vlc_call_back_playing)
        self.vlc_player.add_callback(vlc.EventType.MediaPlayerPaused, self.vlc_call_back_paused)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    # Jet: 这里改写
    # window = ParentWindow()
    window = XunFeiConvertingUIMainWindow()
    '''
    #传递参数到子窗口
    child = ChildWindow(window.main_ui.audio_file, window.main_ui.raw_txt_file)

    # 通过toolButton将两个窗体关联
    btn = window.main_ui.btn_xunfei2txt
    btn.clicked.connect(child.show)  #显示转写状态的子窗口
    '''
    # window.show()  # Jet:有此句会多出现一个窗口！
    # ex = Ui_MainWindow()
    sys.exit(app.exec_())
