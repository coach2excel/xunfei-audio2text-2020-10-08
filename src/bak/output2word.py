import docx
from docx.shared import Inches, Cm

'''
Jet: 
pip install docx  #此法安装，运行会出问题！
命令窗口显示成功后，在开发环境中输入import docx测试三方库有没有成功，显示Import Error: No module named 'exceptions'，查了相关资料

When import docx in python3.3 I have error ImportError: No module named 'exceptions'，

原来是通过命令行下载的docx安装包还没有完全兼容python3，估计这个安装包还只兼容python 2(python2 和python3差别还是挺大的，虽然现在python 3出来很久了，但是不少三方库还没有更新)，因此需要自己找一个兼容的包进行安装，地址为：

https://www.lfd.uci.edu/~gohlke/pythonlibs/

找到python_docx-0.8.10-py2.py3-none-any.whl，点击下载到本地，然后在Anaconda Prompt命令窗口中输入一下指令：

pip uninstall docx （把原来安装的docx卸载掉）
pip install python_docx-0.8.10-py2.py3-none-any.whl
————————————————
版权声明：本文为CSDN博主「miangmiang咩」的原创文章，遵循CC 4.0 BY-SA版权协议，转载请附上原文出处链接及本声明。
原文链接：https://blog.csdn.net/jiangjieqazwsx/article/details/80326214

'''

from docx.oxml.ns import qn  # 字体设置

# 文件名称
file_name = "./test.docx"
# 打开文件
document = docx.Document()

# Jet: 修改正文的中文字体类型，对标题不起作用！
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# level=0：添加“标题”段落，下方会出现一要长细线
document.add_heading('title0： 教练录音转文字', level=0)
# 添加标题(默认一级标题)
document.add_heading('title1 （default)')
# 添加2级标题
document.add_heading('title2', level=2)
# 修改2级标题的字体类型，示例代码：
run = document.add_heading('', level=2).add_run(u"修改2级标题的字体类型为微软雅黑: ")  # 应用场景示例标题
run.font.name = u'微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

# 添加段落
paragraph = document.add_paragraph('passage1: life is short, I like python')
# 使用一个段落作为“光标”，并在其上直接插入一个新段落
prior_paragraph = paragraph.insert_paragraph_before('insert_paragraph_before')

# Jet：这个插入功能很好！灵活性大，可以在之后插入内容

# 添加分页符
document.add_page_break()
# 添加段落
paragraph = document.add_paragraph('passage2 on page 2: life is short, I like python')
# 添加表格的方法
table = document.add_table(rows=3, cols=2)
# 行和列指示是基于零的，就像在列表访问中一样
cell = table.cell(0, 1)
# 在第一个、第二列输入内容
cell.text = 'parrot, possibly dead'
# 选中第二行
row = table.rows[1]
# 分别在第二行的第一列、第二列输入内容
row.cells[0].text = 'speaker1:'
row.cells[1].text = '说话文字字节占比'
row2 = table.rows[2]
row2.cells[0].text = 'speaker2:'
row2.cells[1].text = '说话文字字节占比'

# 在.rows和.columns上的集合是可迭代的，这样你就可以直接在使用它们for循环
for row in table.rows:
    for cell in row.cells:
        print(cell.text)
for column in table.columns:
    for cell in column.cells:
        print(cell.text)
# 在表中的行或列的计数，只要使用len()的顺序
row_count = len(table.rows)
col_count = len(table.columns)
print(row_count)
print(col_count)
# 以递增方式向表中添加行
row = table.add_row()
# Word具有一组预格式化的表格样式
table.style = 'LightShading-Accent1'

# 添加图片
document.add_picture('abc.png')
# 指定其宽度或高度，如英寸，如果仅指定一个，这样的高宽比是保留的
document.add_picture('abc.png', width=Inches(1.0))
# 厘米
document.add_picture('abc.png', width=Cm(5.0))

# 应用段落样式
document.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')
# 等同于下边两个语句
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
paragraph.style = 'ListBullet'

# 应用粗体和斜体
paragraph = document.add_paragraph('测试')
paragraph.add_run('加粗字体').bold = True
# 等同于以下三个语句
run = paragraph.add_run('加粗字体')
run.bold = True
paragraph.add_run('测试.')
# 不是必须提供的文字给.add_paragraph()方法。
# 这可以使你的代码更简单，如果你从建立段从运行反正：
paragraph = document.add_paragraph()
paragraph.add_run('测试')
paragraph.add_run('文字加粗').bold = True
paragraph.add_run(' 的功能.')

# 应用字符样式
paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')  # 斜体
# 等同于以下三个语句
paragraph = document.add_paragraph('Normal text, ')
run = paragraph.add_run('text with emphasis.')
run.style = 'Emphasis'

# 将其保存到名为“test.docx”的文件中
document.save(file_name)
'''
————————————————
版权声明：本文为CSDN博主「梦因you而美」的原创文章，遵循CC 4.0 BY-SA版权协议，转载请附上原文出处链接及本声明。
原文链接：https://blog.csdn.net/apollo_miracle/article/details/88721213
'''
